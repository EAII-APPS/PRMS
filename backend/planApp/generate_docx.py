from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from django.core.files.storage import default_storage
from django.http import JsonResponse
from rest_framework.decorators import api_view, permission_classes
from rest_framework import permissions, status
from rest_framework.response import Response
from .models import StrategicGoal, MainGoal, KPI,PlanDocument
from userApp.models import *
from monApp.models import SystemSetting
from rest_framework import permissions, status
from .docx_data import generate_docx_data,generate_three_data
from datetime import datetime
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import *
from docx.oxml import OxmlElement,ns
from docx.oxml.ns import qn
import os
import tempfile
import shutil
# from docx2pdf import convert
from django.conf import settings
# import pythoncom
import subprocess
from django.conf import settings
from django.db import transaction
import logging
from django.http import FileResponse
# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)



current_year = datetime.now().year



# Function to set "header-like" properties to a run
def set_Title_style(paragraph, text):
    custom_font = "Ebrima"
    Title_font_size = Pt(20)
    run = paragraph.add_run()
    utf8_text = text.encode('utf-8').decode('utf-8')
    run.text = utf8_text
    run.bold = True
    run.font.name = custom_font
    run.font.size = Title_font_size
    run.font.color.rgb = RGBColor(255, 188, 20)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def set_Header1_style(paragraph, text,docx):
    custom_font = "Ebrima"
    Header1_font_size = Pt(17)
    run = paragraph.add_run(text)
    
    # Set font, size, and color
    run.font.name = custom_font
    run.font.size = Header1_font_size
    run.font.color.rgb = RGBColor(47, 110, 188)
    run.bold = True
    
    # Apply Heading 1 style for TOC
    paragraph.style = docx.styles['Heading 1']
def set_celltext_style(run, text):
    custom_font = "Ebrima"
    run.text = text
    run.bold = True
    run.font.name = custom_font
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(47, 110, 188)

def set_Header2_style(paragraph, text,docx):
    custom_font = "Ebrima"
    Header2_font_size = Pt(15)
    run = paragraph.add_run(text)
    
    # Set font, size, and color
    run.font.name = custom_font
    run.font.size = Header2_font_size
    run.font.color.rgb = RGBColor(49, 128, 224)
    run.bold = True
    
    # Apply Heading 2 style for TOC
    paragraph.style = docx.styles['Heading 2']

def Header_style(run, text):
    custom_font = "Ebrima"
    Header2_font_size = Pt(15)
    run.text = text
    run.bold = True
    run.font.name = custom_font
    run.font.size = Header2_font_size

def style_paragraph(paragraph, text, font_name="Ebrima", font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY_LOW):
    """
    Styles the first run of a paragraph, sets paragraph alignment, and adds text if the paragraph is empty.
    Adds a non-breaking space at the end if alignment is set to justify to avoid excessive gaps.
    
    Args:
        paragraph: The paragraph object to style.
        text (str): The text to add if the paragraph is empty.
        font_name (str): The font name to set.
        font_size (int or Pt): The font size to set (default is 12 pt).
        alignment: The paragraph alignment (default is justified).
    """
    # Set alignment for the paragraph
    paragraph.alignment = alignment

    # Add text if paragraph has no runs
    if not paragraph.runs:
        run = paragraph.add_run(text)
    else:
        run = paragraph.runs[0]
    
    # Apply font styling
    run.font.name = font_name
    run.font.size = Pt(font_size) if isinstance(font_size, int) else font_size
    
    # Adjust spacing for justified alignment to avoid uneven gaps
    if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        if not text.endswith("\n"):  # Only add if it doesn't already end with a line break
            paragraph.add_run("\u00A0")  # Add a non-breaking space

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_table_of_contents(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    # This inserts the TOC field
    fld_code = OxmlElement('w:fldSimple')
    fld_code.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    run._r.append(fld_code)

def set_cell_text_style(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Ebrima'


def create_docx_table(document, table_data):
    # Add a table with 1 row and 4 columns (adjust as needed based on your data)
    table = document.add_table(rows=1, cols=4)
    print("generating first table ...")
    # Set alignment of the table to center
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table style
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    Blue = parse_xml(r'<w:shd {} w:fill="110d7c"/>'.format(nsdecls('w')))
    for i in range(4):
        Blue = parse_xml(r'<w:shd {} w:fill="110d7c"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(Blue)
    
    # Set headers for the table
    hdr_cells[0].text = 'ስትራቴጂክ ግቦች'
    hdr_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].text = 'ዋና ዋና ተግባራት'
    hdr_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    hdr_cells[2].text = 'ቁልፍ የአፈፃፀም አመልካቾች (KPIs)'
    hdr_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[2].paragraphs[0].runs[0].font.bold = True
    hdr_cells[3].text = 'የሚጠበቁ ቁልፍ ውጤቶች'
    hdr_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[3].paragraphs[0].runs[0].font.bold = True
    
    # Add strategic goals and main goals
    for strategic_goal in table_data:
        LightBlue = parse_xml(r'<w:shd {} w:fill="D6E3F1"/>'.format(nsdecls('w')))
        
        # Determine the number of rows this strategic goal will span
        strategic_goal_rows = sum(len(main_goal['kpis']) for main_goal in strategic_goal['maingoals']) + len(strategic_goal['maingoals'])
        
        # Create the first row for this strategic goal
        start_row = len(table.rows)
        row_cells = table.add_row().cells
        row_cells[0].text = strategic_goal['strategic_goal_name']
        row_cells[0]._tc.get_or_add_tcPr().append(LightBlue)
        
        main_goal_count = 0
        for main_goal in strategic_goal['maingoals']:
            Gray = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
            # Create the first row for this main goal
            if main_goal_count == 0:
                row_cells[1].text = main_goal['main_goal_name'] if main_goal['main_goal_name'] is not None else "-"
                row_cells[3].text = main_goal['expected_outcome'] if main_goal['expected_outcome'] is not None else "-"
                row_cells[1]._tc.get_or_add_tcPr().append(Gray)
                row_cells[3]._tc.get_or_add_tcPr().append(Gray)
                main_goal_start_row = len(table.rows) - 1
            else:
                row_cells = table.add_row().cells
                row_cells[1].text = main_goal['main_goal_name'] if main_goal['main_goal_name'] is not None else "-"
                row_cells[3].text = main_goal['expected_outcome'] if main_goal['expected_outcome'] is not None else "-"
                row_cells[1]._tc.get_or_add_tcPr().append(Gray)
                row_cells[3]._tc.get_or_add_tcPr().append(Gray)
                main_goal_start_row = len(table.rows) - 1
            


            for idx, kpi in enumerate(main_goal['kpis']):
                kpi_name = kpi['kpi_name'] if kpi['kpi_name'] is not None else "-"

                # For the first KPI, append the text directly to the existing paragraph
                if idx == 0 and main_goal_count == 0:
                    paragraph = row_cells[2].paragraphs[0]
                else:
                    paragraph = row_cells[2].add_paragraph()  # Add a new paragraph for each subsequent KPI

                run = paragraph.add_run(kpi_name)
                paragraph.style = 'ListBullet'  # Apply 'ListBullet' style to each paragraph
            # Merge the main goal and expected outcome cells vertically
            main_goal_end_row = len(table.rows) - 1
            table.cell(main_goal_start_row, 1).merge(table.cell(main_goal_end_row, 1))
            table.cell(main_goal_start_row, 3).merge(table.cell(main_goal_end_row, 3))
            main_goal_count += 1
        
        # Merge the strategic goal cell vertically
        end_row = len(table.rows) - 1
        table.cell(start_row, 0).merge(table.cell(end_row, 0))

def create_plan_table(document,filter_year, table_data):
    # Add a table with 1 row and 11 columns (adjust as needed based on your data)
    print("table data her:",table_data)
    six_month = False
    nine_month = False
    # Iterate through table_data
    for strategic_goal in table_data:
        # Iterate through maingoals
        for main_goal in strategic_goal['maingoals']:
            # Iterate through KPIs
            for kpi in main_goal['kpis']:
                if 'pl6' in kpi:
                    six_month = True
                if 'pl9' in kpi:
                    nine_month = True


    if six_month or nine_month:
        table = document.add_table(rows=1, cols=5)
    else:  
        table = document.add_table(rows=1, cols=8)
    table.width = None
    # Set alignment of the table to center
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # table.style = "Light Grid"
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    Blue = parse_xml(r'<w:shd {} w:fill="110d7c"/>'.format(nsdecls('w')))
    if six_month or nine_month:
        for i in range(5):
            Blue = parse_xml(r'<w:shd {} w:fill="110d7c"/>'.format(nsdecls('w')))
            hdr_cells[i]._tc.get_or_add_tcPr().append(Blue)
    else:
        for i in range(8):
            Blue = parse_xml(r'<w:shd {} w:fill="110d7c"/>'.format(nsdecls('w')))
            hdr_cells[i]._tc.get_or_add_tcPr().append(Blue)
    
    # Set headers for the table
    hdr_cells[0].text = 'ስትራቴጂክ ግቦች፣ ዋና ዋና ተግባራት እና ዝርዝር ቁልፍ የአፈፃፀም አመልካቾች'
    hdr_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].text = 'መለኪያ'
    hdr_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    hdr_cells[2].text = f'{filter_year-1} \n መነሻ'
    hdr_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[2].paragraphs[0].runs[0].font.bold = True
    hdr_cells[3].text = f'{filter_year} \n ግብ'
    hdr_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[3].paragraphs[0].runs[0].font.bold = True

    if six_month:
        hdr_cells[4].text = '6 ወር'
        hdr_cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
        hdr_cells[4].paragraphs[0].runs[0].font.bold = True
    elif nine_month: 
        hdr_cells[4].text = '9 ወር'
        hdr_cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
        hdr_cells[4].paragraphs[0].runs[0].font.bold = True
    else:
        hdr_cells[4].text = '1ኛ'
        hdr_cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
        hdr_cells[4].paragraphs[0].runs[0].font.bold = True
        hdr_cells[5].text = '2ኛ'
        hdr_cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
        hdr_cells[5].paragraphs[0].runs[0].font.bold = True
        hdr_cells[6].text = '3ኛ'
        hdr_cells[6].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
        hdr_cells[6].paragraphs[0].runs[0].font.bold = True
        hdr_cells[7].text = '4ኛ' 
        hdr_cells[7].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
        hdr_cells[7].paragraphs[0].runs[0].font.bold = True
    if six_month or nine_month:
        for i in range(5):
            set_cell_text_style(hdr_cells[i])
    else:
        for i in range(8):
            set_cell_text_style(hdr_cells[i])
    # Iterate over table_data and fill the table cells
    i=0
    for strategic_goal in table_data:
        Orange = parse_xml(r'<w:shd {} w:fill="ffbc14"/>'.format(nsdecls('w')))
        row_cells = table.add_row().cells
        i+=1
        string = f"ስትራቴጂክ ግብ {i}፡-"
        if six_month or nine_month:
            row_cells[0].merge(row_cells[4])  # Merge all cells in the row
        else:
            row_cells[0].merge(row_cells[7])  # Merge all cells in the row
        row_cells[0]._tc.get_or_add_tcPr().append(Orange)
        thetext = strategic_goal['strategic_goal_name'] if strategic_goal['strategic_goal_name'] is not None else "-"
        row_cells[0].text = f'{string} {thetext}'
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(17, 13, 124)  # Blue Color
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_text_style(row_cells[0])
        for main_goal in strategic_goal['maingoals']:
            Gray = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            row_cells = table.add_row().cells
            if six_month or nine_month:
                row_cells[0].merge(row_cells[4])  # Merge all cells in the row
            else:
                row_cells[0].merge(row_cells[7])  # Merge all cells in the row
            row_cells[0]._tc.get_or_add_tcPr().append(Gray)
            row_cells[0].text = main_goal['main_goal_name'] if main_goal['main_goal_name'] is not None else "-"
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            set_cell_text_style(row_cells[0])
            
            for kpi in main_goal['kpis']:
                row_cells = table.add_row().cells
                row_cells[0].text = kpi['kpi_name'] if kpi['kpi_name'] is not None else "-"
                row_cells[1].text = kpi['measure'] if kpi['measure'] else '-'
                row_cells[2].text = str(kpi['initial']) +""+str(kpi.get('initial_unit',"")) if kpi['initial'] is not None else "-" # You need to adjust this based on your data structure
                row_cells[3].text = str(kpi['annual_value'])+""+str(kpi.get('annual_unit',"")) if kpi['annual_value'] is not None else "-" # Adjust as needed
                if 'pl6' in kpi:
                    row_cells[4].text = str(kpi.get('pl6', "-")+" " + str(kpi.get('annual_unit', "")))
                elif 'pl9' in kpi: 
                    row_cells[4].text = str(kpi.get('pl9', "-")) + " " + str(kpi.get('annual_unit', ""))
                else:
                    row_cells[4].text = str(kpi.get('pl1', "-"))+ "" + str(kpi.get('pl1_unit', ""))
                    row_cells[5].text = str(kpi.get('pl2', "-"))+ "" + str(kpi.get('pl2_unit', ""))
                    row_cells[6].text = str(kpi.get('pl3', "-"))+ "" + str(kpi.get('pl3_unit', ""))
                    row_cells[7].text = str(kpi.get('pl4', "-"))+ "" + str(kpi.get('pl4_unit', ""))
                    if six_month or nine_month:
                        for j in range(5):
                            set_cell_text_style(hdr_cells[j])
                    else:
                        for j in range(8):
                            set_cell_text_style(hdr_cells[j])
                    
def create_three_plan_table(document,filter_year, table_data):
    # Add a table with 1 row and 11 columns (adjust as needed based on your data)

    table = document.add_table(rows=1, cols=7)
    table.width = None
    print("generating second table ...")
    # Set alignment of the table to center
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # table.style = "Light Grid"
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    Blue = parse_xml(r'<w:shd {} w:fill="110d7c"/>'.format(nsdecls('w')))
    for i in range(7):
        Blue = parse_xml(r'<w:shd {} w:fill="110d7c"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(Blue)

    # Set headers for the table
    hdr_cells[0].text = 'ስትራቴጂክ ግቦች፣ ዋና ዋና ተግባራት እና ዝርዝር ቁልፍ የአፈፃፀም አመልካቾች'
    hdr_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].text = 'መለኪያ'
    hdr_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    hdr_cells[2].text = f'መነሻ'
    hdr_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[2].paragraphs[0].runs[0].font.bold = True
    hdr_cells[3].text = f'{filter_year}-{filter_year + 2} \n ግብ'
    hdr_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[3].paragraphs[0].runs[0].font.bold = True

    hdr_cells[4].text = f'{filter_year}'
    hdr_cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[4].paragraphs[0].runs[0].font.bold = True
    hdr_cells[5].text = f'{filter_year + 1}'
    hdr_cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[5].paragraphs[0].runs[0].font.bold = True
    hdr_cells[6].text = f'{filter_year + 2}'
    hdr_cells[6].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[6].paragraphs[0].runs[0].font.bold = True
    for i in range(7):
        set_cell_text_style(hdr_cells[i])
    
    # Iterate over table_data and fill the table cells
    i=0
    for strategic_goal in table_data:
        Orange = parse_xml(r'<w:shd {} w:fill="ffbc14"/>'.format(nsdecls('w')))
        row_cells = table.add_row().cells
        i+=1
        string = f"ስትራቴጂክ ግብ {i}፡-"
        row_cells[0].merge(row_cells[6])  # Merge all cells in the row
        row_cells[0]._tc.get_or_add_tcPr().append(Orange)
        thetext = strategic_goal['strategic_goal_name'] if strategic_goal['strategic_goal_name'] is not None else "-"
        row_cells[0].text = f'{string} {thetext}'
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(17, 13, 124)  # Blue Color
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_text_style(row_cells[0])
        for main_goal in strategic_goal['maingoals']:
            Gray = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            row_cells = table.add_row().cells

            row_cells[0].merge(row_cells[6])  # Merge all cells in the row
            row_cells[0]._tc.get_or_add_tcPr().append(Gray)
            row_cells[0].text = main_goal['main_goal_name'] if main_goal['main_goal_name'] is not None else "-"
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            set_cell_text_style(row_cells[0])
            
            for kpi in main_goal['kpis']:
                row_cells = table.add_row().cells
                row_cells[0].text = kpi['kpi_name'] if kpi['kpi_name'] is not None else "-"
                row_cells[1].text = kpi['measure'] if kpi['measure'] else '-'
                row_cells[2].text = f"{kpi.get('initial', '-')} "
                row_cells[3].text = f"{kpi.get('three_year', '-')} {kpi.get('three_year_unit', '')}"
                row_cells[4].text = f"{kpi.get('year_one_value', '-')} "
                row_cells[5].text = f"{kpi.get('year_two_value', '-')} "
                row_cells[6].text = f"{kpi.get('year_three_value', '-')} "
                for j in range(7):
                    set_cell_text_style(row_cells[j])

                    
def create_three_performance_table(document,filter_year, table_data):
    # Add a table with 1 row and 11 columns (adjust as needed based on your data)

    table = document.add_table(rows=1, cols=7)
    table.width = None
    print("generating second table ...")
    # Set alignment of the table to center
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # table.style = "Light Grid"
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    Blue = parse_xml(r'<w:shd {} w:fill="110d7c"/>'.format(nsdecls('w')))
    for i in range(7):
        Blue = parse_xml(r'<w:shd {} w:fill="110d7c"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(Blue)

    # Set headers for the table
    hdr_cells[0].text = 'ስትራቴጂክ ግቦች፣ ዋና ዋና ተግባራት እና ዝርዝር ቁልፍ የአፈፃፀም አመልካቾች'
    hdr_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].text = 'መለኪያ'
    hdr_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    hdr_cells[2].text = f'መነሻ'
    hdr_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[2].paragraphs[0].runs[0].font.bold = True
    hdr_cells[3].text = f'{filter_year}-{filter_year + 2} \n አፈጻጸም'
    hdr_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[3].paragraphs[0].runs[0].font.bold = True

    hdr_cells[4].text = f'{filter_year}'
    hdr_cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[4].paragraphs[0].runs[0].font.bold = True
    hdr_cells[5].text = f'{filter_year + 1}'
    hdr_cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[5].paragraphs[0].runs[0].font.bold = True
    hdr_cells[6].text = f'{filter_year + 2}'
    hdr_cells[6].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 188, 20)  # Orange Color
    hdr_cells[6].paragraphs[0].runs[0].font.bold = True
    for i in range(7):
        set_cell_text_style(hdr_cells[i])

    # Iterate over table_data and fill the table cells
    i=0
    for strategic_goal in table_data:
        Orange = parse_xml(r'<w:shd {} w:fill="ffbc14"/>'.format(nsdecls('w')))
        row_cells = table.add_row().cells
        i+=1
        string = f"ስትራቴጂክ ግብ {i}፡-"
        row_cells[0].merge(row_cells[6])  # Merge all cells in the row
        row_cells[0]._tc.get_or_add_tcPr().append(Orange)
        thetext = strategic_goal['strategic_goal_name'] if strategic_goal['strategic_goal_name'] is not None else "-"
        row_cells[0].text = f'{string} {thetext}'
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(17, 13, 124)  # Blue Color
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_text_style(row_cells[0])
        for main_goal in strategic_goal['maingoals']:
            Gray = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            row_cells = table.add_row().cells

            row_cells[0].merge(row_cells[6])  # Merge all cells in the row
            row_cells[0]._tc.get_or_add_tcPr().append(Gray)
            row_cells[0].text = main_goal['main_goal_name'] if main_goal['main_goal_name'] is not None else "-"
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            set_cell_text_style(row_cells[0])
            
            for kpi in main_goal['kpis']:
                performance_data = kpi.get('performance_data', {})
                row_cells = table.add_row().cells
                row_cells[0].text = kpi['kpi_name'] if kpi['kpi_name'] is not None else "-"
                row_cells[1].text = kpi['measure'] if kpi['measure'] else '-'
                row_cells[2].text = f"{kpi.get('initial', '-')}"
                row_cells[3].text = f"{performance_data.get('total_performance', '-')}"
                row_cells[4].text = f"{performance_data.get('year_one_performance', '-')}"
                row_cells[5].text = f"{performance_data.get('year_two_performance', '-')}"
                row_cells[6].text = f"{performance_data.get('year_three_performance', '-')}"
                for j in range(7):
                    set_cell_text_style(row_cells[j])


@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def generate_docx(request):
    document_id = request.query_params.get('doc_id')
    division = request.query_params.get('division')
    sector = request.query_params.get('sector')
    user = request.user
    
    try:
        # Fetch the PlanDocument object from the database
        plandocument = PlanDocument.objects.prefetch_related('plan_narrations__subtitles').get(id=document_id)
    except PlanDocument.DoesNotExist:
        return Response({"error": "PlanDocument not found"}, status=404)
    filter_year = plandocument.year
    filter_quarter= plandocument.quarter

    if sector or division is not None:
        if sector is not None:
            table_data = generate_docx_data(
                filter_year=filter_year,
                filter_sector=sector,
                filter_division=None,
                filter_quarter=filter_quarter,
                user=request.user,
            )
        elif division is not None:
            document_user=User.objects.filter(id=division)
            table_data = generate_docx_data(
                filter_year=filter_year,
                filter_sector=None,
                filter_division=division,
                filter_quarter=filter_quarter,
                user=request.user,
            )
    if (user.is_superuser or user.monitoring_id):
        document_user=User.objects.get(id=plandocument.added_by.id)
        table_data = generate_docx_data(
            filter_year=filter_year,
            filter_sector=sector,
            filter_division=division,
            filter_quarter=filter_quarter,
            user=document_user,
        )
    else:
        table_data = generate_docx_data(
            filter_year=filter_year,
            filter_sector=sector,
            filter_division=division,
            filter_quarter=filter_quarter,
            user=request.user,
        )
    docx = Document()
    styles = docx.styles
    custom_font = "Ebrima"
    Paragraph_style = styles['Normal']
    Paragraph_style.font.name = 'Ebrima'
    Paragraph_style.font.size = Pt(12)
    Title_font_size = Pt(20)
    Header1_font_size = Pt(17)
    Header2_font_size = Pt(15)
    Paragraph_font_size = Pt(12)

    section = docx.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    



    six_month = False
    nine_month = False
    first_quarter = False
    second_quarter = False
    third_quarter = False
    fourth_quarter = False
    # Iterate through table_data
    if filter_quarter:
        if filter_quarter == 1:
            first_quarter = True
        elif filter_quarter == 2:
            second_quarter = True
        elif filter_quarter == 3:
            third_quarter = True
        elif filter_quarter == 4:
            fourth_quarter = True
        elif filter_quarter == 6:
            six_month = True
        elif filter_quarter == 9:
            nine_month = True

    if six_month:
        header_para.text = f"የኢትዮጵያ አርቲፊሻል ኢንተለጀንስ ኢንስቲትዩት የ{plandocument.year} በጀት ዓመት የ ስድስት ወር የልማት ዕቅድ"
    elif nine_month:
        header_para.text = f"የኢትዮጵያ አርቲፊሻል ኢንተለጀንስ ኢንስቲትዩት የ{plandocument.year} በጀት ዓመት የ ዘጠኝ ወር የልማት ዕቅድ"
    elif first_quarter:
        header_para.text = f"የኢትዮጵያ አርቲፊሻል ኢንተለጀንስ ኢንስቲትዩት የ{plandocument.year} በጀት ዓመት የ 1ኛ ሩብ ዓመት ዕቅድ"
    elif second_quarter:
        header_para.text = f"የኢትዮጵያ አርቲፊሻል ኢንተለጀንስ ኢንስቲትዩት የ{plandocument.year} በጀት ዓመት የ 2ኛ ሩብ ዓመት ዕቅድ"
    elif third_quarter:
        header_para.text = f"የኢትዮጵያ አርቲፊሻል ኢንተለጀንስ ኢንስቲትዩት የ{plandocument.year} በጀት ዓመት የ 3ኛ ሩብ ዓመት ዕቅድ"
    elif fourth_quarter:
        header_para.text = f"የኢትዮጵያ አርቲፊሻል ኢንተለጀንስ ኢንስቲትዩት የ{plandocument.year} በጀት ዓመት የ 4ኛ ሩብ ዓመት ዕቅድ"
    else:
        header_para.text = f"የኢትዮጵያ አርቲፊሻል ኢንተለጀንስ ኢንስቲትዩት የ{plandocument.year} በጀት ዓመት ዓመታዊ የልማት ዕቅድ"

    system_setting = SystemSetting.objects.first()
    print("writing document ...")
    if system_setting is not None:
        logo_image_url = system_setting.logo_image
        section = docx.sections[0]
        section.top_margin = Inches(1)  # Adjust as needed
        section.bottom_margin = Inches(1)  # Adjust as needed
        # Add an image and center it
        image_paragraph = docx.add_paragraph()
        image_run = image_paragraph.add_run()
        image_run.add_picture(logo_image_url, width=Inches(2))
        image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Access the footer
    footer = docx.sections[0].footer
    # Add the first paragraph for the date
    date_paragraph = footer.add_paragraph()
    Fottertext = date_paragraph.add_run("AI for All ! & አርቲፊሻል ኢንተለጀንስ ለሁሉም !")

    # Set alignment, font size, and color for the date
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    Fottertext.font.size = Pt(10)
    Fottertext.font.bold = False
    # Fottertext.font.color.rgb = RGBColor(255, 188, 20)

    title_paragraph = docx.add_paragraph()
    title_text = str(plandocument.title)  # Ensure it's a string
    set_Title_style(title_paragraph, title_text)



    # Set the section to have a different header/footer for the first page
    section.different_first_page_header_footer = True

    # Access the footer for the first page only
    footer = section.first_page_footer

    # Add the first paragraph for the date
    date_paragraph = footer.add_paragraph()
    date_run = date_paragraph.add_run(f"{plandocument.year} ዓ.ም")

    # Set alignment, font size, and color for the date
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run.font.size = Pt(20)
    date_run.font.bold = True
    date_run.font.color.rgb = RGBColor(17, 13, 124)

    # Add another paragraph for the location
    location_paragraph = footer.add_paragraph()
    location_run = location_paragraph.add_run("አዲስ አበባ")
    location_run.font.bold = True

    # Set alignment, font size, and color for the location
    location_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    location_run.font.size = Pt(20)
    location_run.font.color.rgb = RGBColor(17, 13, 124)
    # Add a page break
    docx.add_page_break()

    toc_paragraph = docx.add_paragraph()
    toc_paragraph.add_run("Table of Contents").bold = True
    add_table_of_contents(docx)
    # Add a page break or space after the TOC placeholder
    docx.add_page_break()
    add_page_number(docx.sections[0].footer.paragraphs[0].add_run())
    # First Section
    for index, plan_narration in enumerate(plandocument.plan_narrations.filter(section_type="መግቢያ"), start=1):
        # Add section title
        title_paragraph = docx.add_paragraph()
        set_Header1_style(title_paragraph, f'{index}. መግቢያ',docx)
        paragraph = docx.add_paragraph()
        style_paragraph(paragraph, text=plan_narration.description)
        # Iterate through photos in the current subtitle and add them to the document
        plan_photo = plan_narration.Plan_photos.all()
        if plan_photo:
            for plan_photo in plan_narration.Plan_photos.all():
                if plan_photo.photos:
                    photo_path = plan_photo.photos.path
                    if os.path.exists(photo_path):
                        try:
                            docx.add_picture(photo_path, width=Inches(1.5))
                        except Exception as e:
                            logger.error(f"Failed to add picture {photo_path}: {e}")
        i = 1
        # Iterate through subtitles of the current PlanNarration
        for i, subtitle in enumerate(plan_narration.subtitles.all(), start=1):
            # Add header for subtitle description
            header_paragraph = docx.add_paragraph()
            set_Header2_style(header_paragraph, f'{index}.{i} {subtitle.subtitle}',docx)
            paragraph = docx.add_paragraph()
            style_paragraph(paragraph, text=subtitle.description)
            # Iterate through photos in the current subtitle and add them to the document
            plan_photo = subtitle.Plan_photos.all()
            if plan_photo:
                for plan_photo in subtitle.Plan_photos.all():
                    if plan_photo.photos:
                        photo_path = plan_photo.photos.path
                        if os.path.exists(photo_path):
                            try:
                                docx.add_picture(photo_path, width=Inches(1.5))
                            except Exception as e:
                                logger.error(f"Failed to add picture {photo_path}: {e}")

        # After iterating through subtitles, add the additional subtitle
        i += 1  # Increment the index for the new subtitle
        header_paragraph = docx.add_paragraph()
        set_Header2_style(header_paragraph, f'{index}.{i} ስትራቴጂክ ግቦች',docx)
        strategic_goals = StrategicGoal.objects.filter(is_deleted=0)
        for goal in strategic_goals:
            docx.add_paragraph(goal.name, style='ListNumber')


    #Second Section
    for index, plan_narration in enumerate(plandocument.plan_narrations.filter(section_type="የልማት ዕቅዱ እንድምታዎች እና የሚጠበቁ ውጤቶች"), start=2):
        # Add section title
        title_paragraph = docx.add_paragraph()
        set_Header1_style(title_paragraph, f'{index}. {plan_narration.title}',docx)
        paragraph = docx.add_paragraph()
        style_paragraph(paragraph, text=plan_narration.description)
        plan_photo =  plan_narration.Plan_photos.all()
        if plan_photo :
            for plan_photo in plan_narration.Plan_photos.all():
                if plan_photo.photos:
                    photo_path = plan_photo.photos.path
                    if os.path.exists(photo_path):
                        try:
                            docx.add_picture(photo_path, width=Inches(1.5))
                        except Exception as e:
                            logger.error(f"Failed to add picture {photo_path}: {e}")
        # Iterate through subtitles of the current PlanNarration
        for i, subtitle in enumerate(plan_narration.subtitles.all(), start=1):
            # Add header for subtitle description
            header_paragraph = docx.add_paragraph()
            set_Header2_style(header_paragraph, f'{index}.{i} {subtitle.subtitle}',docx)
            paragraph = docx.add_paragraph()
            style_paragraph(paragraph, text=subtitle.description)
            # Iterate through photos in the current subtitle and add them to the document
            plan_photo = subtitle.Plan_photos.all()
            if plan_photo :
                for plan_photo in subtitle.Plan_photos.all():
                    if plan_photo.photos:
                        photo_path = plan_photo.photos.path
                        if os.path.exists(photo_path):
                            try:
                                docx.add_picture(photo_path, width=Inches(1.5))
                            except Exception as e:
                                logger.error(f"Failed to add picture {photo_path}: {e}")

    header_paragraph = docx.add_paragraph()
    Header_style(header_paragraph.add_run(), f'ስትራቴጂካዊ ግቦች፣ ዋና ዋና ተግባራት፣ ቁልፍ አመላካቾች እና የሚጠበቁ ቁልፍ ውጤቶች')
    # Add a table 
    create_docx_table(docx,table_data)

    # Add a page break
    docx.add_page_break()
     #Third Section
    for index, plan_narration in enumerate(plandocument.plan_narrations.filter(section_type="የበጀት ዓመት ኢላማዎችና የድርጊት መርሃ ግብር"), start=3):
        # Add section title
        title_paragraph = docx.add_paragraph()
        set_Header1_style(title_paragraph, f'{index}. {plan_narration.title}',docx)
        paragraph = docx.add_paragraph()
        style_paragraph(paragraph, text=plan_narration.description)
        plan_photo = plan_narration.Plan_photos.all()
        if plan_photo :
            for plan_photo in plan_narration.Plan_photos.all():
                if plan_photo.photos:
                    photo_path = plan_photo.photos.path
                    if os.path.exists(photo_path):
                        try:
                            docx.add_picture(photo_path, width=Inches(1.5))
                        except Exception as e:
                            logger.error(f"Failed to add picture {photo_path}: {e}")
        # Iterate through subtitles of the current PlanNarration
        for i, subtitle in enumerate(plan_narration.subtitles.all(), start=1):
            # Add header for subtitle description
            header_paragraph = docx.add_paragraph()
            set_Header2_style(header_paragraph, f'{index}.{i} {subtitle.subtitle}',docx)
            paragraph = docx.add_paragraph()
            style_paragraph(paragraph, text=subtitle.description)
                # Iterate through photos in the current subtitle and add them to the document
            plan_photo = subtitle.Plan_photos.all()
            if plan_photo :    
                for plan_photo in subtitle.Plan_photos.all():
                    if plan_photo.photos:
                        photo_path = plan_photo.photos.path
                        if os.path.exists(photo_path):
                            try:
                                docx.add_picture(photo_path, width=Inches(1.5))
                            except Exception as e:
                                logger.error(f"Failed to add picture {photo_path}: {e}")
    create_plan_table(docx,filter_year,table_data)
     #Fourth Section
    for index, plan_narration in enumerate(plandocument.plan_narrations.filter(section_type="የበጀት ዓመቱን የልማት ዕቅድ ለማሰካት የሚተገበሩ ማስፈጸሚያዎች"), start=4):
        # Add section title
        title_paragraph = docx.add_paragraph()
        set_Header1_style(title_paragraph, f'{index}. {plan_narration.title}',docx)
        paragraph = docx.add_paragraph()
        style_paragraph(paragraph, text=plan_narration.description)
        plan_photo = plan_narration.Plan_photos.all()
        if plan_photo:
            for plan_photo in plan_narration.Plan_photos.all():
                if plan_photo.photos:
                    photo_path = plan_photo.photos.path
                    if os.path.exists(photo_path):
                        try:
                            docx.add_picture(photo_path, width=Inches(1.5))
                        except Exception as e:
                            logger.error(f"Failed to add picture {photo_path}: {e}")
        # Iterate through subtitles of the current PlanNarration
        for i, subtitle in enumerate(plan_narration.subtitles.all(), start=1):
            # Add header for subtitle description
            header_paragraph = docx.add_paragraph()
            set_Header2_style(header_paragraph, f'{index}.{i} {subtitle.subtitle}',docx)
            paragraph = docx.add_paragraph()
            style_paragraph(paragraph, text=subtitle.description)
            # Iterate through photos in the current subtitle and add them to the document
            plan_photo = subtitle.Plan_photos.all()
            if plan_photo:
                for plan_photo in subtitle.Plan_photos.all():
                    if plan_photo.photos:
                        photo_path = plan_photo.photos.path
                        if os.path.exists(photo_path):
                            try:
                                docx.add_picture(photo_path, width=Inches(1.5))
                            except Exception as e:
                                logger.error(f"Failed to add picture {photo_path}: {e}")
    
    try:
        # Initialize COM library
        # pythoncom.CoInitialize()

        # Prepare file paths
        docx_file_name = f"{plandocument.title.replace(' ', '_')}.docx"
        final_docx_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", docx_file_name)
        final_pdf_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", docx_file_name.replace(".docx", ".pdf"))

        os.makedirs(os.path.dirname(final_docx_path), exist_ok=True)

        # Create and save the DOCX file directly
        docx.save(final_docx_path)

        # Convert the DOCX file to PDF
        # convert(final_docx_path, final_pdf_path)
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", os.path.dirname(final_pdf_path),
            final_docx_path
        ], check=True)

        # Prepare response
        response = {
            'docx_file_path': f"http://127.0.0.1:8000/media/plandocuments/{docx_file_name}",
            'pdf_file_path': f"http://127.0.0.1:8000/media/plandocuments/{docx_file_name.replace('.docx', '.pdf')}"
        }

        print(f"Document opened from {os.path.basename(final_pdf_path)}")
        return JsonResponse(response, status=200)
    
    except subprocess.CalledProcessError as e:
        logger.error(f"LibreOffice conversion failed: {e}")
        return JsonResponse({'error': 'Failed to convert DOCX to PDF'}, status=500)


    except Exception as e:
        # Enhanced error logging
        print(f"An error occurred: {e}")
        return JsonResponse({'error': str(e)}, status=500)

    finally:
        # Ensure COM library is uninitialized
        # pythoncom.CoUninitialize()
        print("Document generation process completed.")

@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def generate_pdf(request):

    document = Document()
    # Set page orientation to landscape (swap width and height)
    sections = document.sections
    for section in sections:
        # Swap width and height for landscape orientation
        section.page_width = Inches(11)  # Landscape width (standard letter width)
        section.page_height = Inches(8.5)  # Landscape height (standard letter height)

    year = request.query_params.get('year')
    division = request.query_params.get('division')
    sector = request.query_params.get('sector')
    quarter = request.query_params.get('quarter')
    user = request.user

    if quarter is None:
        table_data = generate_docx_data(
            filter_year=year,
            filter_sector=sector,
            filter_division=division,
            filter_quarter= 12,
            user=user,
        )
    else:
        table_data = generate_docx_data(
            filter_year=year,
            filter_sector=sector,
            filter_division=division,
            filter_quarter= int(quarter),
            user=user,
        )

    title_paragraph = document.add_paragraph()
    set_Title_style(title_paragraph,f"የ {year} የእቅድ ሠንጠረዥ")

    create_plan_table(document,int(year),table_data)
    # Define the file paths
    docx_file_name = "Plan_table.docx"
    pdf_file_name = docx_file_name.replace(".docx", ".pdf")
    docx_file_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", docx_file_name)
    pdf_file_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", pdf_file_name)
    # pythoncom.CoInitialize()
    # Ensure the directories exist
    os.makedirs(os.path.dirname(docx_file_path), exist_ok=True)
    
    # Save the DOCX file
    document.save(docx_file_path)
    
    # Convert the DOCX file to PDF
    # convert(docx_file_path, pdf_file_path)
    subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", os.path.dirname(pdf_file_path),
            docx_file_path
        ], check=True)
    
    # Check if the PDF file was created and respond accordingly
    if not os.path.exists(pdf_file_path):
        return JsonResponse({'error': 'PDF conversion failed'}, status=500)
    
        # Check if the PDF file was created successfully
    if not os.path.exists(pdf_file_path):
        return JsonResponse({'error': 'PDF conversion failed'}, status=500)

    # Open the PDF file in binary read mode and send it in the response
    return FileResponse(open(pdf_file_path, 'rb'), content_type='application/pdf', as_attachment=True, filename=pdf_file_name)
    
@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def generate_word(request):

    document = Document()
    # Set page orientation to landscape (swap width and height)
    sections = document.sections
    for section in sections:
        # Swap width and height for landscape orientation
        section.page_width = Inches(11)  # Landscape width (standard letter width)
        section.page_height = Inches(8.5)  # Landscape height (standard letter height)

    year = request.query_params.get('year')
    division = request.query_params.get('division')
    sector = request.query_params.get('sector')
    quarter = request.query_params.get('quarter')
    user = request.user

    if quarter is None:
        table_data = generate_docx_data(
            filter_year=year,
            filter_sector=sector,
            filter_division=division,
            filter_quarter= 12,
            user=user,
        )
    else:
        table_data = generate_docx_data(
            filter_year=year,
            filter_sector=sector,
            filter_division=division,
            filter_quarter= int(quarter),
            user=user,
        )
    title_paragraph = document.add_paragraph()
    set_Title_style(title_paragraph,f"የ {year} የእቅድ ሠንጠረዥ")

    create_plan_table(document,int(year),table_data)
    # Define the file paths
    docx_file_name = "Plan_table.docx"
    # pdf_file_name = docx_file_name.replace(".docx", ".pdf")
    docx_file_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", docx_file_name)
    # pdf_file_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", pdf_file_name)
    # pythoncom.CoInitialize()
    # Ensure the directories exist
    os.makedirs(os.path.dirname(docx_file_path), exist_ok=True)
    
    # Save the DOCX file
    document.save(docx_file_path)
    
    # Open the PDF file in binary read mode and send it in the response
    return FileResponse(open(docx_file_path, 'rb'),   content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, filename=docx_file_name) 

@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def generate_word_three(request):

    document = Document()
    # Set page orientation to landscape (swap width and height)
    sections = document.sections
    for section in sections:
        # Swap width and height for landscape orientation
        section.page_width = Inches(11)  # Landscape width (standard letter width)
        section.page_height = Inches(8.5)  # Landscape height (standard letter height)

    year = request.query_params.get('year')
    division = request.query_params.get('division')
    sector = request.query_params.get('sector')
    quarter = request.query_params.get('quarter')
    user = request.user

    if quarter is None:
        table_data = generate_three_data(
            filter_year=year,
            filter_sector=sector,
            filter_division=division,
            user=user,
        )
    else:
        table_data = generate_three_data(
            filter_year=year,
            filter_sector=sector,
            filter_division=division,
            user=user,
        )
    title_paragraph = document.add_paragraph()
    set_Title_style(title_paragraph, f"የ {year}-{str(int(year) + 2)} የእቅድ ሠንጠረዥ")

    create_three_plan_table(document,int(year),table_data)
    # Define the file paths
    docx_file_name = "Three_Plan_table.docx"
    # pdf_file_name = docx_file_name.replace(".docx", ".pdf")
    docx_file_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", docx_file_name)
    # pdf_file_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", pdf_file_name)
    # pythoncom.CoInitialize()
    # Ensure the directories exist
    os.makedirs(os.path.dirname(docx_file_path), exist_ok=True)
    
    # Save the DOCX file
    document.save(docx_file_path)
    
    # Open the PDF file in binary read mode and send it in the response
    return FileResponse(open(docx_file_path, 'rb'),   content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, filename=docx_file_name) 

@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def generate_pdf_three(request):

    document = Document()
    # Set page orientation to landscape (swap width and height)
    sections = document.sections
    for section in sections:
        # Swap width and height for landscape orientation
        section.page_width = Inches(11)  # Landscape width (standard letter width)
        section.page_height = Inches(8.5)  # Landscape height (standard letter height)

    year = request.query_params.get('year')
    division = request.query_params.get('division')
    sector = request.query_params.get('sector')
    quarter = request.query_params.get('quarter')
    user = request.user

    if quarter is None:
        table_data = generate_three_data(
            filter_year=year,
            filter_sector=sector,
            filter_division=division,
            user=user,
        )
    else:
        table_data = generate_three_data(
            filter_year=year,
            filter_sector=sector,
            filter_division=division,
            user=user,
        )
    title_paragraph = document.add_paragraph()
    set_Title_style(title_paragraph, f"የ {year}-{str(int(year) + 2)} የእቅድ ሠንጠረዥ")

    create_three_plan_table(document,int(year),table_data)
       # Define the file paths
    docx_file_name = "Plan_table.docx"
    pdf_file_name = docx_file_name.replace(".docx", ".pdf")
    docx_file_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", docx_file_name)
    pdf_file_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", pdf_file_name)
    # pythoncom.CoInitialize()
    # Ensure the directories exist
    os.makedirs(os.path.dirname(docx_file_path), exist_ok=True)
    
    # Save the DOCX file
    document.save(docx_file_path)
    
    # Convert the DOCX file to PDF
    # convert(docx_file_path, pdf_file_path)
    subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", os.path.dirname(pdf_file_path),
            docx_file_path
        ], check=True)

    # Check if the PDF file was created and respond accordingly
    if not os.path.exists(pdf_file_path):
        return JsonResponse({'error': 'PDF conversion failed'}, status=500)
    
        # Check if the PDF file was created successfully
    if not os.path.exists(pdf_file_path):
        return JsonResponse({'error': 'PDF conversion failed'}, status=500)

    # Open the PDF file in binary read mode and send it in the response
    return FileResponse(open(pdf_file_path, 'rb'), content_type='application/pdf', as_attachment=True, filename=pdf_file_name)
    
@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def generate_word_three_performance(request):

    document = Document()
    # Set page orientation to landscape (swap width and height)
    sections = document.sections
    for section in sections:
        # Swap width and height for landscape orientation
        section.page_width = Inches(11)  # Landscape width (standard letter width)
        section.page_height = Inches(8.5)  # Landscape height (standard letter height)

    year = request.query_params.get('year')
    division = request.query_params.get('division')
    sector = request.query_params.get('sector')
    quarter = request.query_params.get('quarter')
    user = request.user

    if quarter is None:
        table_data = generate_three_data(
            filter_year=year,
            filter_sector=sector,
            filter_division=division,
            user=user,
        )
    else:
        table_data = generate_three_data(
            filter_year=year,
            filter_sector=sector,
            filter_division=division,
            user=user,
        )
    title_paragraph = document.add_paragraph()
    set_Title_style(title_paragraph, f"የ {year}-{str(int(year) + 2)} የአፈጻጸም ሠንጠረዥ")

    create_three_performance_table(document,int(year),table_data)
    # Define the file paths
    docx_file_name = "Three_Plan_table.docx"
    # pdf_file_name = docx_file_name.replace(".docx", ".pdf")
    docx_file_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", docx_file_name)
    # pdf_file_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", pdf_file_name)
    # pythoncom.CoInitialize()
    # Ensure the directories exist
    os.makedirs(os.path.dirname(docx_file_path), exist_ok=True)
    
    # Save the DOCX file
    document.save(docx_file_path)
    
    # Open the PDF file in binary read mode and send it in the response
    return FileResponse(open(docx_file_path, 'rb'),   content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, filename=docx_file_name) 

@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def generate_pdf_three_performance(request):

    document = Document()
    # Set page orientation to landscape (swap width and height)
    sections = document.sections
    for section in sections:
        # Swap width and height for landscape orientation
        section.page_width = Inches(11)  # Landscape width (standard letter width)
        section.page_height = Inches(8.5)  # Landscape height (standard letter height)

    year = request.query_params.get('year')
    division = request.query_params.get('division')
    sector = request.query_params.get('sector')
    quarter = request.query_params.get('quarter')
    user = request.user

    if quarter is None:
        table_data = generate_three_data(
            filter_year=year,
            filter_sector=sector,
            filter_division=division,
            user=user,
        )
    else:
        table_data = generate_three_data(
            filter_year=year,
            filter_sector=sector,
            filter_division=division,
            user=user,
        )
    title_paragraph = document.add_paragraph()
    set_Title_style(title_paragraph, f"የ {year}-{str(int(year) + 2)} የአፈጻጸም ሠንጠረዥ")

    create_three_performance_table(document,int(year),table_data)
       # Define the file paths
    docx_file_name = "Plan_table.docx"
    pdf_file_name = docx_file_name.replace(".docx", ".pdf")
    docx_file_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", docx_file_name)
    pdf_file_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", pdf_file_name)
    # pythoncom.CoInitialize()
    # Ensure the directories exist
    os.makedirs(os.path.dirname(docx_file_path), exist_ok=True)
    
    # Save the DOCX file
    document.save(docx_file_path)
    
    # Convert the DOCX file to PDF
    # convert(docx_file_path, pdf_file_path)
    subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", os.path.dirname(pdf_file_path),
            docx_file_path
        ], check=True)

    # Check if the PDF file was created and respond accordingly
    if not os.path.exists(pdf_file_path):
        return JsonResponse({'error': 'PDF conversion failed'}, status=500)
    
        # Check if the PDF file was created successfully
    if not os.path.exists(pdf_file_path):
        return JsonResponse({'error': 'PDF conversion failed'}, status=500)

    # Open the PDF file in binary read mode and send it in the response
    return FileResponse(open(pdf_file_path, 'rb'), content_type='application/pdf', as_attachment=True, filename=pdf_file_name)
   

@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def merge_doc(request):
    document_ids = request.query_params.getlist('documentIds')
    division = request.query_params.get('division')
    sector = request.query_params.get('sector')

    if not document_ids:
        return Response({"error": "No document IDs provided"}, status=400)
    
    # Create a new Document object for merging
    docx = Document()
    styles = docx.styles
    # Iterate through each document ID provided
    for doc_id in document_ids:
        try:
            # Fetch the document from the database
            plandocument = PlanDocument.objects.prefetch_related('plan_narrations__subtitles').get(id=doc_id)  # Adjust with your actual model and field
            # document_path = plandocument.file.path  # Adjust if needed based on your file field
            user = request.user
            filter_year = plandocument.year
            filter_quarter= plandocument.quarter

            if sector or division is not None:
                if sector is not None:
                    table_data = generate_docx_data(
                        filter_year=filter_year,
                        filter_sector=sector,
                        filter_division=None,
                        filter_quarter=filter_quarter,
                        user=request.user,
                    )
                elif division is not None:
                    document_user=User.objects.filter(id=division)
                    table_data = generate_docx_data(
                        filter_year=filter_year,
                        filter_sector=None,
                        filter_division=division,
                        filter_quarter=filter_quarter,
                        user=request.user,
                    )
            if (user.is_superuser or user.monitoring_id):
                document_user=User.objects.get(email=plandocument.added_by)
                table_data = generate_docx_data(
                    filter_year=filter_year,
                    filter_sector=sector,
                    filter_division=division,
                    filter_quarter=filter_quarter,
                    user=document_user,
                )
            else:
                table_data = generate_docx_data(
                    filter_year=filter_year,
                    filter_sector=sector,
                    filter_division=division,
                    filter_quarter=filter_quarter,
                    user=request.user,
                )
            custom_font = "Ebrima"
            Paragraph_style = styles['Normal']
            Paragraph_style.font.name = 'Ebrima'
            Paragraph_style.font.size = Pt(12)
            Title_font_size = Pt(20)
            Header1_font_size = Pt(17)
            Header2_font_size = Pt(15)
            Paragraph_font_size = Pt(12)

            system_setting = SystemSetting.objects.first()
            print("writing document ...")
            if system_setting is not None:
                logo_image_url = system_setting.logo_image
                section = docx.sections[0]
                section.top_margin = Inches(2)  # Adjust as needed
                section.bottom_margin = Inches(2)  # Adjust as needed
                # Add an image and center it
                image_paragraph = docx.add_paragraph()
                image_run = image_paragraph.add_run()
                image_run.add_picture(logo_image_url, width=Inches(2))
                image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            title_paragraph = docx.add_paragraph()
            set_Title_style(title_paragraph,f"{plandocument.title}")

            docx.add_page_break()
            add_page_number(docx.sections[0].footer.paragraphs[0].add_run())

            # First Section
            for index, plan_narration in enumerate(plandocument.plan_narrations.filter(section_type="መግቢያ"), start=1):
                # Add section title
                title_paragraph = docx.add_paragraph()
                set_Header1_style(title_paragraph, f'{index}. {plan_narration.title}',docx)
                paragraph = docx.add_paragraph()
                style_paragraph(paragraph, text=plan_narration.description)
                # Iterate through photos in the current subtitle and add them to the document
                plan_photo = plan_narration.Plan_photos.all()
                if plan_photo:
                    for plan_photo in plan_narration.Plan_photos.all():
                        if plan_photo.photos:
                            photo_path = plan_photo.photos.path
                            if os.path.exists(photo_path):
                                try:
                                    docx.add_picture(photo_path, width=Inches(1.5))
                                except Exception as e:
                                    logger.error(f"Failed to add picture {photo_path}: {e}")
                i = 1
                # Iterate through subtitles of the current PlanNarration
                for i, subtitle in enumerate(plan_narration.subtitles.all(), start=1):
                    # Add header for subtitle description
                    header_paragraph = docx.add_paragraph()
                    set_Header2_style(header_paragraph, f'{index}.{i} {subtitle.subtitle}',docx)
                    paragraph = docx.add_paragraph()
                    style_paragraph(paragraph, text=subtitle.description)
                    # Iterate through photos in the current subtitle and add them to the document
                    plan_photo = subtitle.Plan_photos.all()
                    if plan_photo:
                        for plan_photo in subtitle.Plan_photos.all():
                            if plan_photo.photos:
                                photo_path = plan_photo.photos.path
                                if os.path.exists(photo_path):
                                    try:
                                        docx.add_picture(photo_path, width=Inches(1.5))
                                    except Exception as e:
                                        logger.error(f"Failed to add picture {photo_path}: {e}")

                # After iterating through subtitles, add the additional subtitle
                i += 1  # Increment the index for the new subtitle
                header_paragraph = docx.add_paragraph()
                set_Header2_style(header_paragraph, f'{index}.{i} ስትራቴጂክ ግቦች',docx)
                strategic_goals = StrategicGoal.objects.filter(is_deleted=0)
                for goal in strategic_goals:
                    docx.add_paragraph(goal.name, style='ListNumber')


            #Second Section
            for index, plan_narration in enumerate(plandocument.plan_narrations.filter(section_type="የልማት ዕቅዱ እንድምታዎች እና የሚጠበቁ ውጤቶች"), start=2):
                # Add section title
                title_paragraph = docx.add_paragraph()
                set_Header1_style(title_paragraph, f'{index}. {plan_narration.title}',docx)
                paragraph = docx.add_paragraph()
                style_paragraph(paragraph, text=plan_narration.description)
                plan_photo =  plan_narration.Plan_photos.all()
                if plan_photo :
                    for plan_photo in plan_narration.Plan_photos.all():
                        if plan_photo.photos:
                            photo_path = plan_photo.photos.path
                            if os.path.exists(photo_path):
                                try:
                                    docx.add_picture(photo_path, width=Inches(1.5))
                                except Exception as e:
                                    logger.error(f"Failed to add picture {photo_path}: {e}")
                # Iterate through subtitles of the current PlanNarration
                for i, subtitle in enumerate(plan_narration.subtitles.all(), start=1):
                    # Add header for subtitle description
                    header_paragraph = docx.add_paragraph()
                    set_Header2_style(header_paragraph, f'{index}.{i} {subtitle.subtitle}',docx)
                    paragraph = docx.add_paragraph()
                    style_paragraph(paragraph, text=subtitle.description)
                    # Iterate through photos in the current subtitle and add them to the document
                    plan_photo = subtitle.Plan_photos.all()
                    if plan_photo :
                        for plan_photo in subtitle.Plan_photos.all():
                            if plan_photo.photos:
                                photo_path = plan_photo.photos.path
                                if os.path.exists(photo_path):
                                    try:
                                        docx.add_picture(photo_path, width=Inches(1.5))
                                    except Exception as e:
                                        logger.error(f"Failed to add picture {photo_path}: {e}")

            header_paragraph = docx.add_paragraph()
            Header_style(header_paragraph.add_run(), f'ስትራቴጂካዊ ግቦች፣ ዋና ዋና ተግባራት፣ ቁልፍ አመላካቾች እና የሚጠበቁ ቁልፍ ውጤቶች')
            # Add a table 
            create_docx_table(docx,table_data)

            # Add a page break
            docx.add_page_break()
            #Third Section
            for index, plan_narration in enumerate(plandocument.plan_narrations.filter(section_type="የበጀት ዓመት ኢላማዎችና የድርጊት መርሃ ግብር"), start=3):
                # Add section title
                title_paragraph = docx.add_paragraph()
                set_Header1_style(title_paragraph, f'{index}. {plan_narration.title}',docx)
                paragraph = docx.add_paragraph()
                style_paragraph(paragraph, text=plan_narration.description)
                plan_photo = plan_narration.Plan_photos.all()
                if plan_photo:
                    for plan_photo in plan_narration.Plan_photos.all():
                        if plan_photo.photos:
                            photo_path = plan_photo.photos.path
                            if os.path.exists(photo_path):
                                try:
                                    docx.add_picture(photo_path, width=Inches(1.5))
                                except Exception as e:
                                    logger.error(f"Failed to add picture {photo_path}: {e}")
                # Iterate through subtitles of the current PlanNarration
                for i, subtitle in enumerate(plan_narration.subtitles.all(), start=1):
                    # Add header for subtitle description
                    header_paragraph = docx.add_paragraph()
                    set_Header2_style(header_paragraph, f'{index}.{i} {subtitle.subtitle}',docx)
                    paragraph = docx.add_paragraph()
                    style_paragraph(paragraph, text=subtitle.description)
                        # Iterate through photos in the current subtitle and add them to the document
                    plan_photo = subtitle.Plan_photos.all()
                    if plan_photo :    
                        for plan_photo in subtitle.Plan_photos.all():
                            if plan_photo.photos:
                                photo_path = plan_photo.photos.path
                                if os.path.exists(photo_path):
                                    try:
                                        docx.add_picture(photo_path, width=Inches(1.5))
                                    except Exception as e:
                                        logger.error(f"Failed to add picture {photo_path}: {e}")
            create_plan_table(docx,filter_year,table_data)
            #Fourth Section
            for index, plan_narration in enumerate(plandocument.plan_narrations.filter(section_type="የበጀት ዓመቱን የልማት ዕቅድ ለማሰካት የሚተገበሩ ማስፈጸሚያዎች"), start=4):
                # Add section title
                title_paragraph = docx.add_paragraph()
                set_Header1_style(title_paragraph, f'{index}. {plan_narration.title}',docx)
                paragraph = docx.add_paragraph()
                style_paragraph(paragraph, text=plan_narration.description)
                plan_photo = plan_narration.Plan_photos.all()
                if plan_photo:
                    for plan_photo in plan_narration.Plan_photos.all():
                        if plan_photo.photos:
                            photo_path = plan_photo.photos.path
                            if os.path.exists(photo_path):
                                try:
                                    docx.add_picture(photo_path, width=Inches(1.5))
                                except Exception as e:
                                    logger.error(f"Failed to add picture {photo_path}: {e}")
                # Iterate through subtitles of the current PlanNarration
                for i, subtitle in enumerate(plan_narration.subtitles.all(), start=1):
                    # Add header for subtitle description
                    header_paragraph = docx.add_paragraph()
                    set_Header2_style(header_paragraph, f'{index}.{i} {subtitle.subtitle}',docx)
                    paragraph = docx.add_paragraph()
                    style_paragraph(paragraph, text=subtitle.description)
                    # Iterate through photos in the current subtitle and add them to the document
                    plan_photo = subtitle.Plan_photos.all()
                    if plan_photo:
                        for plan_photo in subtitle.Plan_photos.all():
                            if plan_photo.photos:
                                photo_path = plan_photo.photos.path
                                if os.path.exists(photo_path):
                                    try:
                                        docx.add_picture(photo_path, width=Inches(1.5))
                                    except Exception as e:
                                        logger.error(f"Failed to add picture {photo_path}: {e}")
            

        except PlanDocument.DoesNotExist:
            return Response({"error": f"Document with ID {doc_id} does not exist."}, status=404)
        except Exception as e:
            return Response({"error": str(e)}, status=500)
    try:
        # Initialize COM library
        # pythoncom.CoInitialize()

        # Prepare file paths
        docx_file_name = f"{plandocument.title.replace(' ', '_')}.docx"
        final_docx_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", docx_file_name)
        final_pdf_path = os.path.join(settings.MEDIA_ROOT, "plandocuments", docx_file_name.replace(".docx", ".pdf"))

        os.makedirs(os.path.dirname(final_docx_path), exist_ok=True)

        # Create and save the DOCX file directly
        docx.save(final_docx_path)

        # Convert the DOCX file to PDF
        # convert(final_docx_path, final_pdf_path)
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", os.path.dirname(final_pdf_path),
            final_docx_path
        ], check=True)

        # Prepare response
        response = {
            'docx_file_path': f"http://127.0.0.1:8000/media/plandocuments/{docx_file_name}",
            'pdf_file_path': f"http://127.0.0.1:8000/media/plandocuments/{docx_file_name.replace('.docx', '.pdf')}"
        }

        print(f"Document opened from {os.path.basename(final_pdf_path)}")
        return JsonResponse(response, status=200)

    except Exception as e:
        # Enhanced error logging
        print(f"An error occurred: {e}")
        return JsonResponse({'error': str(e)}, status=500)

    finally:
        # Ensure COM library is uninitialized
        # pythoncom.CoUninitialize()
        print("Document generation process completed.")

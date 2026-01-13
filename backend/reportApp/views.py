# from urllib import request
from django.shortcuts import get_object_or_404, render
import docx
from rest_framework import generics, status
from .models import *
from .serializer import *
from rest_framework.response import Response
from rest_framework import permissions, status
from rest_framework.permissions import IsAuthenticated
from django.contrib.auth import get_user_model
import subprocess
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from django.http import JsonResponse
from planApp.models import *
from userApp.models import *

from django.db.models import OuterRef, Subquery
import re
from docx import Document
from docx.shared import Pt, RGBColor
from django.core.files.storage import default_storage
from django.conf import settings
import os
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from django.http import JsonResponse
import tempfile
from docx.shared import Inches
from docx2pdf import convert
# import pythoncom
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from monApp.models import SystemSetting
from django.db.models import Q,Prefetch
from rest_framework.decorators import api_view
from docx.shared import RGBColor
from planApp.unit_calculator import convert_quarterly_plans_to_base,convert_quarterly_performance_to_base
from collections import defaultdict
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def delete_old_documents():
    documents_dir = os.path.join(settings.MEDIA_ROOT, 'documents')
    for root, dirs, files in os.walk(documents_dir):
        for file in files:
            try:
                print(f"Deleting file: {file}")
                file_path = os.path.join(root, file)
                os.remove(file_path)
            except Exception as e:
                print(f"Could not delete {file_path}: {e}")

class MergedAnnualKPIInstance:
    def __init__(self, data):
        self.kpi_id = data['kpi']
        self.kpi = KPI.objects.filter(id=data['kpi']).first()  # Fetch KPI instance
        self.measure_id = data['measure']
        self.measure = Measure.objects.filter(id=data['measure']).first()  # Fetch Measure instance
        self.year = data['year']
        self.annual = data['annual']
        self.initial = data['initial']
        self.pl1 = data['pl1']
        self.pl2 = data['pl2']
        self.pl3 = data['pl3']
        self.pl4 = data['pl4']
        self.pr1 = data['pr1']
        self.pr2 = data['pr2']
        self.pr3 = data['pr3']
        self.pr4 = data['pr4']
        self.weight = data['weight']
        self.operation = data['operation']
        self.annual_unit_id = data['annual_unit_id']
        self.initial_unit_id = data['initial_unit_id']
        self.pl1_unit_id = data['pl1_unit_id']
        self.pl2_unit_id = data['pl2_unit_id']
        self.pl3_unit_id = data['pl3_unit_id']
        self.pl4_unit_id = data['pl4_unit_id']
        self.pr1_unit_id = data['pr1_unit_id']
        self.pr2_unit_id = data['pr2_unit_id']
        self.pr3_unit_id = data['pr3_unit_id']
        self.pr4_unit_id = data['pr4_unit_id']
    
    @property
    def annual_unit(self):
        return Unit.objects.filter(id=self.annual_unit_id).first()  # Returns the Unit instance

    @property
    def initial_unit(self):
        return Unit.objects.filter(id=self.initial_unit_id).first()  # Returns the Unit instance

    @property
    def pl1_unit(self):
        return Unit.objects.filter(id=self.pl1_unit_id).first()  # Returns the Unit instance

    @property
    def pl2_unit(self):
        return Unit.objects.filter(id=self.pl2_unit_id).first()  # Returns the Unit instance

    @property
    def pl3_unit(self):
        return Unit.objects.filter(id=self.pl3_unit_id).first()  # Returns the Unit instance

    @property
    def pl4_unit(self):
        return Unit.objects.filter(id=self.pl4_unit_id).first()  # Returns the Unit instance

    @property
    def pr1_unit(self):
        return Unit.objects.filter(id=self.pr1_unit_id).first()  # Returns the Unit instance

    @property
    def pr2_unit(self):
        return Unit.objects.filter(id=self.pr2_unit_id).first()  # Returns the Unit instance

    @property
    def pr3_unit(self):
        return Unit.objects.filter(id=self.pr3_unit_id).first()  # Returns the Unit instance

    @property
    def pr4_unit(self):
        return Unit.objects.filter(id=self.pr4_unit_id).first()  # Returns the Unit instance


def merge_annual_kpis(annual_kpis):
    # Group by kpi, division_id, measure, year, and the additional unit IDs
    grouped_kpis = defaultdict(lambda: {
        'kpi': None,
        'measure': None,
        'year': None,
        'annual': 0,
        'initial': 0,
        'pl1': 0,
        'pl2': 0,
        'pl3': 0,
        'pl4': 0,
        'pr1': 0,
        'pr2': 0,
        'pr3': 0,
        'pr4': 0,
        'weight': 0,
        'operation': None,
        'pl1_unit_id': None,  # Add placeholder for unit IDs to be set after merging
        'pl2_unit_id': None,
        'pl3_unit_id': None,
        'pl4_unit_id': None,
        'annual_unit_id': None,  # For base unit assignment
        'initial_unit_id': 0,
        'count': 0  # Initialize count for averaging
    })

    for record in annual_kpis:
        converted_plans, base_unit_id = convert_quarterly_plans_to_base(record)
        converted_performance, performance_base_unit_id = convert_quarterly_performance_to_base(record)
        print(f"Converted Plans: {converted_plans}, Base Unit ID: {base_unit_id}")
        print(f"Converted Performance: {converted_performance}, Performance Base Unit ID: {performance_base_unit_id}")
        # Construct key for grouping, now includes sector_id derived from division_id
        key = (
            record.kpi_id,
            record.measure_id,
            record.year
        )

        # Add values to the existing group
        grouped_kpis[key]['kpi'] = record.kpi_id
        grouped_kpis[key]['measure'] = record.measure_id
        grouped_kpis[key]['year'] = record.year
        # Add unit IDs
        grouped_kpis[key]['pl1_unit_id'] = Unit.objects.get(id=base_unit_id)
        grouped_kpis[key]['pl2_unit_id'] = Unit.objects.get(id=base_unit_id)
        grouped_kpis[key]['pl3_unit_id'] = Unit.objects.get(id=base_unit_id)
        grouped_kpis[key]['pl4_unit_id'] = Unit.objects.get(id=base_unit_id)
        grouped_kpis[key]['pr1_unit_id'] = Unit.objects.get(id=performance_base_unit_id)
        grouped_kpis[key]['pr2_unit_id'] = Unit.objects.get(id=performance_base_unit_id)
        grouped_kpis[key]['pr3_unit_id'] = Unit.objects.get(id=performance_base_unit_id)
        grouped_kpis[key]['pr4_unit_id'] = Unit.objects.get(id=performance_base_unit_id)
        # Add the annual and initial values and their units
        grouped_kpis[key]['annual_unit_id'] = base_unit_id
        grouped_kpis[key]['initial_unit_id'] = Unit.objects.get(id=base_unit_id)
        grouped_kpis[key]['weight'] += record.weight or 0
        grouped_kpis[key]['operation'] = record.operation
        # Get the operation type
        operation = record.operation
        if operation == 'sum':
            # Merge values (including the converted quarterly plans)
            grouped_kpis[key]['annual'] += record.annual or 0
            grouped_kpis[key]['initial'] += record.initial or 0
            grouped_kpis[key]['pl1'] += converted_plans['pl1'] # Use converted value for the plans
            grouped_kpis[key]['pl2'] += converted_plans['pl2']
            grouped_kpis[key]['pl3'] += converted_plans['pl3']
            grouped_kpis[key]['pl4'] += converted_plans['pl4']
            grouped_kpis[key]['pr1'] += converted_performance['pr1'] if converted_performance['pr1'] is not None else 0
            grouped_kpis[key]['pr2'] += converted_performance['pr2'] if converted_performance['pr2'] is not None else 0
            grouped_kpis[key]['pr3'] += converted_performance['pr3'] if converted_performance['pr3'] is not None else 0
            grouped_kpis[key]['pr4'] += converted_performance['pr4'] if converted_performance['pr4'] is not None else 0
        elif operation == 'average':
            # Increment the count for averaging
            grouped_kpis[key]['count'] += 1

            # Average the values (including the converted quarterly plans)
            grouped_kpis[key]['annual'] = (
                (grouped_kpis[key]['annual'] * (grouped_kpis[key]['count'] - 1)) + (record.annual or 0)
            ) / grouped_kpis[key]['count']

            grouped_kpis[key]['initial'] = (
                (grouped_kpis[key]['initial'] * (grouped_kpis[key]['count'] - 1)) + (record.initial or 0)
            ) / grouped_kpis[key]['count']

            grouped_kpis[key]['pl1'] = (
                (grouped_kpis[key]['pl1'] * (grouped_kpis[key]['count'] - 1)) + converted_plans['pl1']
            ) / grouped_kpis[key]['count']

            grouped_kpis[key]['pl2'] = (
                (grouped_kpis[key]['pl2'] * (grouped_kpis[key]['count'] - 1)) + converted_plans['pl2']
            ) / grouped_kpis[key]['count']

            grouped_kpis[key]['pl3'] = (
                (grouped_kpis[key]['pl3'] * (grouped_kpis[key]['count'] - 1)) + converted_plans['pl3']
            ) / grouped_kpis[key]['count']

            grouped_kpis[key]['pl4'] = (
                (grouped_kpis[key]['pl4'] * (grouped_kpis[key]['count'] - 1)) + converted_plans['pl4']
            ) / grouped_kpis[key]['count']

            grouped_kpis[key]['pr1'] = (
                (grouped_kpis[key]['pr1'] * (grouped_kpis[key]['count'] - 1)) + (record.pr1 or 0)
            ) / grouped_kpis[key]['count']

            grouped_kpis[key]['pr2'] = (
                (grouped_kpis[key]['pr2'] * (grouped_kpis[key]['count'] - 1)) + (record.pr2 or 0)
            ) / grouped_kpis[key]['count']

            grouped_kpis[key]['pr3'] = (
                (grouped_kpis[key]['pr3'] * (grouped_kpis[key]['count'] - 1)) + (record.pr3 or 0)
            ) / grouped_kpis[key]['count']

            grouped_kpis[key]['pr4'] = (
                (grouped_kpis[key]['pr4'] * (grouped_kpis[key]['count'] - 1)) + (record.pr4 or 0)
            ) / grouped_kpis[key]['count']

    # Now convert grouped_kpis to instances of MergedAnnualKPI (or equivalent class)
    merged_data = [MergedAnnualKPIInstance(kpi) for kpi in grouped_kpis.values()]

    return merged_data


User = get_user_model()

class KPIDescriptionListCreate(generics.ListCreateAPIView):
    
    queryset = KPIDescription.objects.all()
    serializer_class = KPIDescriptionSerializer
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        division_id = request.user.division_id
        sector_id = request.user.sector_id
        monitoring_id = request.user.monitoring_id
        is_superadmin = request.user.is_superadmin
        if division_id:
            kpi_ids = AnnualKPI.objects.filter(division_id=division_id).distinct().values_list('id', flat=True)
            try:
                if 'pk' in kwargs:  # Check if specific instance is requested
                    kpi_description = KPIDescription.objects.get(pk=kwargs['pk'], kpi_id__in=kpi_ids)
                    serializer = self.get_serializer(kpi_description)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else:
                    kpi_descriptions = self.get_queryset().filter(kpi_id__in=kpi_ids)
                    serializer = self.get_serializer(kpi_descriptions, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except KPIDescription.DoesNotExist:
                return Response({"error": "KPIDescription not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        elif sector_id:
            main_goal_ids = MainGoal.objects.filter(sector_id=sector_id).distinct().values_list('id', flat=True)
            kpi_ids = AnnualKPI.objects.filter(main_goal_id__in=main_goal_ids).distinct().values_list('id', flat=True)
            try:
                if 'pk' in kwargs:  # Check if specific instance is requested
                    kpi_description = KPIDescription.objects.get(pk=kwargs['pk'], kpi_id__in=kpi_ids)
                    serializer = self.get_serializer(kpi_description)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else:
                    kpi_descriptions = self.get_queryset().filter(kpi_id__in=kpi_ids)
                    serializer = self.get_serializer(kpi_descriptions, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except KPIDescription.DoesNotExist:
                return Response({"error": "KPIDescription not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
        elif monitoring_id:
            try:
                if 'pk' in kwargs:  # Check if specific instance is requested
                    kpi_description = KPIDescription.objects.get(pk=kwargs['pk'])
                    serializer = self.get_serializer(kpi_description)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else:
                    kpi_descriptions = self.get_queryset()
                    serializer = self.get_serializer(kpi_descriptions, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except KPIDescription.DoesNotExist:
                return Response({"error": "KPIDescription not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        elif is_superadmin:
            try:
                if 'pk' in kwargs:  # Check if specific instance is requested
                    kpi_description = KPIDescription.objects.get(pk=kwargs['pk'])
                    serializer = self.get_serializer(kpi_description)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else:
                    kpi_descriptions = self.get_queryset()
                    serializer = self.get_serializer(kpi_descriptions, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except KPIDescription.DoesNotExist:
                return Response({"error": "KPIDescription not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def post(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        try:
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class KPIDescriptionRetrieveUpdateDelete(generics.RetrieveUpdateDestroyAPIView):
    queryset = KPIDescription.objects.all()
    serializer_class = KPIDescriptionSerializer
    permission_classes = [IsAuthenticated]

    def get_kpidescription(self, pk):
        return get_object_or_404(KPIDescription, pk=pk)
    
    def delete(self, request, pk):
        try:
            kpides = self.get_kpidescription(pk)
            kpides.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # if division_id == kpides.division_id:  # Check if the user is allowed to delete the car
        #     kpides.delete()
        #     return Response(status=status.HTTP_204_NO_CONTENT)
        # else:
        #     return Response(status=status.HTTP_403_FORBIDDEN)

    
    def put(self, request, pk, *args, **kwargs):
        instance = self.get_kpidescription(pk)
        serializer = self.get_serializer(instance, data=request.data, partial=True)
        try:
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)




class MeasureListCreate(generics.ListCreateAPIView):
    queryset = Measure.objects.all()
    serializer_class = MeasureSerializer
    permission_classes = [IsAuthenticated]

    def get_object(self, pk):
        try:
            return Measure.objects.get(pk=pk)
        except Measure.DoesNotExist:
            return Response({"error": "Measure not found"}, status=status.HTTP_404_NOT_FOUND)

    def get(self, request, *args, **kwargs):
        try:
            if 'pk' in kwargs:  # Check if specific instance is requested
                measures = self.get_object(kwargs['pk'])
                serializer = self.get_serializer(measures)
                return Response(serializer.data, status=status.HTTP_200_OK)
            else:
                measures = self.get_queryset()
                serializer = self.get_serializer(measures, many=True)
                return Response(serializer.data, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def post(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        try:
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



class MeasureRetrieveUpdateDelete(generics.RetrieveUpdateDestroyAPIView):
    queryset = Measure.objects.all()
    serializer_class = MeasureSerializer
    permission_classes = [IsAuthenticated]

    def get_object(self, pk):
        return get_object_or_404(Measure, pk=pk)
    
    def delete(self, request, pk):
        try:
            measure = self.get_object(pk)
            measure.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    
    def put(self, request, pk, *args, **kwargs):
        instance = self.get_object(pk)
        serializer = self.get_serializer(instance, data=request.data, partial=True)
        try:
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class UnitListCreate(generics.ListCreateAPIView):
    queryset = Unit.objects.all()
    serializer_class = UnitSerializer
    permission_classes = [IsAuthenticated]

    def get_object(self, pk):
        try:
            return Unit.objects.get(pk=pk)
        except Unit.DoesNotExist:
            return Response({"error": "Unit not found"}, status=status.HTTP_404_NOT_FOUND)

    def get(self, request, *args, **kwargs):
        try:
            if 'pk' in kwargs:  # Check if specific instance is requested
                units = self.get_object(kwargs['pk'])
                serializer = self.get_serializer(units)
                return Response(serializer.data, status=status.HTTP_200_OK)
            else:
                units = self.get_queryset()
                serializer = self.get_serializer(units, many=True)
                return Response(serializer.data, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



class UnitRetrieveUpdateDelete(generics.RetrieveUpdateDestroyAPIView):
    queryset = Unit.objects.all()
    serializer_class = UnitSerializer
    permission_classes = [IsAuthenticated]

    def get_object(self, pk):
        return get_object_or_404(Unit, pk=pk)
    
    def delete(self, request, pk):
        try:
            units = self.get_object(pk)
            units.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    
    def put(self, request, pk, *args, **kwargs):
        instance = self.get_object(pk)
        serializer = self.get_serializer(instance, data=request.data, partial=True)
        try:
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



class SubtitleFilesListCreateView(generics.ListCreateAPIView):
    queryset = SubtitleFiles.objects.all()
    serializer_class = SubtitleFilesSerializer
    permission_classes = [IsAuthenticated]

class SubtitleFilesRetrieveUpdateDestroyView(generics.RetrieveUpdateDestroyAPIView):
    queryset = SubtitleFiles.objects.all()
    serializer_class = SubtitleFilesSerializer
    permission_classes = [IsAuthenticated]

class SummarySubtitleListCreateView(generics.ListCreateAPIView):
    queryset = SummarySubtitle.objects.all()
    serializer_class = SummarySubtitleSerializer
    permission_classes = [IsAuthenticated]

class SummarySubtitleRetrieveUpdateDestroyView(generics.RetrieveUpdateDestroyAPIView):
    queryset = SummarySubtitle.objects.all()
    serializer_class = SummarySubtitleSerializer
    permission_classes = [IsAuthenticated]

    queryset = Summary.objects.all()
    serializer_class = SummarySerializer
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        # Unified filtering: respect user association, then apply query-params
        division_id = getattr(request.user, 'division_id', None)
        sector_id = getattr(request.user, 'sector_id', None)
        monitoring_id = getattr(request.user, 'monitoring_id', None)
        is_superadmin = getattr(request.user, 'is_superadmin', False)

        filter_year = request.query_params.get('year')
        filter_quarter = request.query_params.get('quarter')
        filter_division = request.query_params.get('division')
        filter_sector = request.query_params.get('sector')

        try:
            qs = Summary.objects.all()

            # Apply user-level constraints first (users can't see beyond their association)
            if division_id:
                qs = qs.filter(division_id=division_id)
            elif sector_id:
                qs = qs.filter(sector_id=sector_id)
            elif monitoring_id:
                qs = qs.filter(monitoring_id=monitoring_id)
            # superadmin: no base constraint

            # Apply explicit query params (only if provided). For non-superadmin users these will further narrow results.
            if filter_division:
                try:
                    qs = qs.filter(division_id=int(filter_division))
                except (ValueError, TypeError):
                    pass

            if filter_sector:
                try:
                    qs = qs.filter(sector_id=int(filter_sector))
                except (ValueError, TypeError):
                    pass

            if filter_year:
                try:
                    qs = qs.filter(year=int(filter_year))
                except (ValueError, TypeError):
                    pass

            # Quarter may come as numeric (1,2,3,4) or as string like 'first'
            if filter_quarter:
                q = str(filter_quarter).strip().lower()
                quarter_map = {
                    '1': 'first', '2': 'second', '3': 'third', '4': 'fourth',
                    'first': 'first', 'second': 'second', 'third': 'third', 'fourth': 'fourth',
                    '6': 'six', '9': 'nine', 'year': 'year'
                }
                mapped = quarter_map.get(q)
                if mapped:
                    qs = qs.filter(quarter__iexact=mapped)

            # If pk provided, return single instance from the filtered qs
            if 'pk' in kwargs:
                try:
                    obj = qs.get(pk=kwargs['pk'])
                    serializer = SummarySerializer(obj)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                except Summary.DoesNotExist:
                    return Response({"error": "Summary not found"}, status=status.HTTP_404_NOT_FOUND)

            # Otherwise return list
            serializer = SummarySerializer(qs, many=True)
            return Response(serializer.data, status=status.HTTP_200_OK)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def post(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        try:
            if serializer.is_valid():
                user = request.user
                user_has_required_association = False
                
                if (hasattr(user, 'sector_id') and user.sector_id is not None and hasattr(user, 'division_id') and user.division_id is None):
                    user_has_required_association = True
                    serializer.validated_data['sector_id'] = user.sector_id
                if hasattr(user, 'monitoring_id') and user.monitoring_id is not None:
                    user_has_required_association = True
                    serializer.validated_data['monitoring_id'] = user.monitoring_id
                if hasattr(user, 'division_id') and user.division_id is not None:
                    user_has_required_association = True
                    serializer.validated_data['division_id'] = user.division_id

                if user_has_required_association:
                    serializer.save()
                    return Response(serializer.data, status=status.HTTP_201_CREATED)
                else:
                    return Response({"error": "User must belong to at least one sector_id, monitoring_id, or division_id."}, status=status.HTTP_400_BAD_REQUEST)

            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
class SummaryListCreate(generics.ListCreateAPIView):
    queryset = Summary.objects.all()
    serializer_class = SummarySerializer
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        division_id = getattr(request.user, 'division_id', None)
        sector_id = getattr(request.user, 'sector_id', None)
        monitoring_id = getattr(request.user, 'monitoring_id', None)
        is_superadmin = getattr(request.user, 'is_superadmin', False)
        filter_year = request.query_params.get('year')
        filter_quarter = request.query_params.get('quarter')
        filter_division = request.query_params.get('division')
        filter_sector = request.query_params.get('sector')

        try:
            qs = Summary.objects.all()
            if division_id:
                qs = qs.filter(division_id=division_id)
            elif sector_id:
                from userApp.models import Division
                sector_divisions = Division.objects.filter(sector_id=sector_id).values_list('id', flat=True)
                from django.db.models import Q
                qs = qs.filter(
                    Q(sector_id=sector_id) |
                    Q(division_id__in=sector_divisions)
                )
            elif monitoring_id:
                pass
            elif is_superadmin:
                pass
                
            else:
                qs = qs.none()

            if filter_division:
                try:
                    div_id = int(filter_division)
                    if division_id:
                        if div_id == division_id.id:
                            qs = qs.filter(division_id=div_id)
                        else:
                            return Response({"error": "Division users can only access their own division data"}, 
                                          status=status.HTTP_403_FORBIDDEN)
                    elif sector_id:
                        from userApp.models import Division
                        allowed_divisions = Division.objects.filter(
                            sector_id=sector_id, 
                            id=div_id
                        ).exists()
                        if allowed_divisions:
                            qs = qs.filter(division_id=div_id)
                        else:
                            return Response({"error": "Sector users can only filter by divisions in their sector"}, 
                                          status=status.HTTP_403_FORBIDDEN)
                    else:
                        qs = qs.filter(division_id=div_id)
                        
                except (ValueError, TypeError):
                    pass

            if filter_sector:
                try:
                    sec_id = int(filter_sector)
                    if monitoring_id or is_superadmin:
                        qs = qs.filter(sector_id=sec_id)
                    elif sector_id:
                        if sec_id == sector_id.id:
                            qs = qs.filter(sector_id=sec_id)
                        else:
                            return Response({"error": "Sector users can only access their own sector data"}, 
                                          status=status.HTTP_403_FORBIDDEN)
                    elif division_id:
                        return Response({"error": "Division users cannot filter by sector"}, 
                                      status=status.HTTP_403_FORBIDDEN)
                        
                except (ValueError, TypeError):
                    pass

            if filter_year:
                try:
                    qs = qs.filter(year=int(filter_year))
                except (ValueError, TypeError):
                    pass

            if filter_quarter:
                q = str(filter_quarter).strip().lower()
                quarter_map = {
                    '1': 'first', '2': 'second', '3': 'third', '4': 'fourth',
                    'first': 'first', 'second': 'second', 'third': 'third', 'fourth': 'fourth',
                    '6': 'six', '9': 'nine', 'year': 'year'
                }
                mapped = quarter_map.get(q)
                if mapped:
                    qs = qs.filter(quarter__iexact=mapped)

            if 'pk' in kwargs:
                try:
                    obj = qs.get(pk=kwargs['pk'])
                    serializer = SummarySerializer(obj)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                except Summary.DoesNotExist:
                    return Response({"error": "Summary not found"}, status=status.HTTP_404_NOT_FOUND)

            serializer = SummarySerializer(qs, many=True)
            return Response(serializer.data, status=status.HTTP_200_OK)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def post(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        try:
            if serializer.is_valid():
                user = request.user
                user_has_required_association = False
                
                if (hasattr(user, 'sector_id') and user.sector_id is not None and hasattr(user, 'division_id') and user.division_id is None):
                    user_has_required_association = True
                    serializer.validated_data['sector_id'] = user.sector_id
                if hasattr(user, 'monitoring_id') and user.monitoring_id is not None:
                    user_has_required_association = True
                    serializer.validated_data['monitoring_id'] = user.monitoring_id
                if hasattr(user, 'division_id') and user.division_id is not None:
                    user_has_required_association = True
                    serializer.validated_data['division_id'] = user.division_id

                if user_has_required_association:
                    serializer.save()
                    return Response(serializer.data, status=status.HTTP_201_CREATED)
                else:
                    return Response({"error": "User must belong to at least one sector_id, monitoring_id, or division_id."}, status=status.HTTP_400_BAD_REQUEST)

            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class SummaryRetrieveUpdateDelete(generics.RetrieveUpdateDestroyAPIView):
    queryset = Summary.objects.all()
    serializer_class = SummarySerializer
    permission_classes = [IsAuthenticated]

    def get_summary(self, pk):
        return get_object_or_404(Summary, pk=pk)
    
    def delete(self, request, pk):
        try:
            summary = self.get_summary(pk)
            summary.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def put(self, request, pk, *args, **kwargs):
        instance = self.get_summary(pk)
        serializer = self.get_serializer(instance, data=request.data, partial=True)
        try:
            if serializer.is_valid():
                user = request.user
                user_has_required_association = False
                
                if (hasattr(user, 'sector_id') and user.sector_id is not None and hasattr(user, 'division_id') and user.division_id is None):
                    user_has_required_association = True
                    serializer.validated_data['sector_id'] = user.sector_id
                if hasattr(user, 'monitoring_id') and user.monitoring_id is not None:
                    user_has_required_association = True
                    serializer.validated_data['monitoring_id'] = user.monitoring_id
                if hasattr(user, 'division_id') and user.division_id is not None:
                    user_has_required_association = True
                    serializer.validated_data['division_id'] = user.division_id

                if user_has_required_association:
                    serializer.save()
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else:
                    return Response({"error": "User must belong to at least one sector_id, monitoring_id, or division_id."}, status=status.HTTP_400_BAD_REQUEST)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



class SummaryFilesListCreate(generics.ListCreateAPIView):
    queryset = SummaryFiles.objects.all()
    serializer_class = SummaryFilesSerializer
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        division_id = request.user.division_id
        sector_id = request.user.sector_id
        monitoring_id = request.user.monitoring_id
        is_superadmin = request.user.is_superadmin
        if division_id:
            sum_id = Summary.objects.filter(division_id = division_id).values_list('id', flat=True)
            try:
                if 'pk' in kwargs: 
                    summary = SummaryFiles.objects.get(pk=kwargs['pk'], summary__in=sum_id)
                    serializer = SummaryFilesSerializer(summary)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else: 
                    summaries = SummaryFiles.objects.filter(summary__in=sum_id)
                    serializer = SummaryFilesSerializer(summaries, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except SummaryFiles.DoesNotExist:
                return Response({"error": "Summary Photo not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        elif sector_id:
            sum_id = Summary.objects.filter(sector_id = sector_id).values_list('id', flat=True)
            try:
                if 'pk' in kwargs: 
                    summary = SummaryFiles.objects.get(pk=kwargs['pk'], summary__in = sum_id)
                    serializer = SummaryFilesSerializer(summary)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else: 
                    summaries = SummaryFiles.objects.filter(summary__in=sum_id)
                    serializer = SummaryFilesSerializer(summaries, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except SummaryFiles.DoesNotExist:
                return Response({"error": "Summary Photo not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
        elif monitoring_id:
            sum_id = Summary.objects.filter(monitoring_id = monitoring_id).values_list('id', flat=True)
            try:
                if 'pk' in kwargs: 
                    summary = SummaryFiles.objects.get(pk=kwargs['pk'], summary__in=sum_id)
                    serializer = SummaryFilesSerializer(summary)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else: 
                    summaries = SummaryFiles.objects.filter(summary__in=sum_id)
                    serializer = SummaryFilesSerializer(summaries, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except SummaryFiles.DoesNotExist:
                return Response({"error": "Summary Photo not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
        elif is_superadmin:
            try:
                if 'pk' in kwargs: 
                    summary = SummaryFiles.objects.get(pk=kwargs['pk'])
                    serializer = SummaryFilesSerializer(summary)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else: 
                    summaries = SummaryFiles.objects.all()
                    serializer = SummaryFilesSerializer(summaries, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except SummaryFiles.DoesNotExist:
                return Response({"error": "Summary Photo not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def post(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        try:
            if serializer.is_valid():
                user = request.user
                user_has_required_association = False
                
                if (hasattr(user, 'sector_id') and user.sector_id is not None and hasattr(user, 'division_id') and user.division_id is None):
                    user_has_required_association = True
                    serializer.validated_data['sector_id'] = user.sector_id
                if hasattr(user, 'monitoring_id') and user.monitoring_id is not None:
                    user_has_required_association = True
                    serializer.validated_data['monitoring_id'] = user.monitoring_id
                if hasattr(user, 'division_id') and user.division_id is not None:
                    user_has_required_association = True
                    serializer.validated_data['division_id'] = user.division_id

                if user_has_required_association:
                    serializer.save()
                    return Response(serializer.data, status=status.HTTP_201_CREATED)
                else:
                    return Response({"error": "User must belong to at least one sector_id, monitoring_id, or division_id."}, status=status.HTTP_400_BAD_REQUEST)

            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
class SummaryFilesRetrieveUpdateDelete(generics.RetrieveUpdateDestroyAPIView):
    queryset = SummaryFiles.objects.all()
    serializer_class = SummaryFilesSerializer
    permission_classes = [IsAuthenticated]

    def get_summaryfiles(self, pk):
        return get_object_or_404(SummaryFiles, pk=pk)
    
    def delete(self, request, pk):
        try:
            summaryfiles = self.get_summaryfiles(pk)
            summaryfiles.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


    def put(self, request, pk, *args, **kwargs):
        instance = self.get_summaryfiles(pk)
        serializer = self.get_serializer(instance, data=request.data, partial=True)
        try:
            if serializer.is_valid():
                user = request.user
                user_has_required_association = False
                
                if (hasattr(user, 'sector_id') and user.sector_id is not None and hasattr(user, 'division_id') and user.division_id is None):
                    user_has_required_association = True
                    serializer.validated_data['sector_id'] = user.sector_id
                if hasattr(user, 'monitoring_id') and user.monitoring_id is not None:
                    user_has_required_association = True
                    serializer.validated_data['monitoring_id'] = user.monitoring_id
                if hasattr(user, 'division_id') and user.division_id is not None:
                    user_has_required_association = True
                    serializer.validated_data['division_id'] = user.division_id

                if user_has_required_association:
                    serializer.save()
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else:
                    return Response({"error": "User must belong to at least one sector_id, monitoring_id, or division_id."}, status=status.HTTP_400_BAD_REQUEST)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class KPIPhotosListCreate(generics.ListCreateAPIView):
    
    queryset = KPIPhotos.objects.all()
    serializer_class = KPIPhotosSerializer
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        division_id = request.user.division_id
        sector_id = request.user.sector_id
        monitoring_id = request.user.monitoring_id
        is_superadmin = request.user.is_superadmin
        if division_id:
            kpi_ids = KPI.objects.filter(division_id=division_id).distinct().values_list('id', flat=True)
            try:
                if 'pk' in kwargs:  # Check if specific instance is requested
                    kpi_photos = KPIPhotos.objects.get(pk=kwargs['pk'], kpi_id__in=kpi_ids)
                    serializer = self.get_serializer(kpi_photos)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else:
                    kpi_photos = self.get_queryset().filter(kpi_id__in=kpi_ids)
                    serializer = self.get_serializer(kpi_photos, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except KPIPhotos.DoesNotExist:
                return Response({"error": "KPI Photo not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        elif sector_id:
            main_goal_ids = MainGoal.objects.filter(sector_id=sector_id).distinct().values_list('id', flat=True)
            kpi_ids = KPI.objects.filter(main_goal_id__in=main_goal_ids).distinct().values_list('id', flat=True)
            try:
                if 'pk' in kwargs:  # Check if specific instance is requested
                    kpi_photos = KPIPhotos.objects.get(pk=kwargs['pk'], kpi_id__in=kpi_ids)
                    serializer = self.get_serializer(kpi_photos)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else:
                    kpi_photos = self.get_queryset().filter(kpi_id__in=kpi_ids)
                    serializer = self.get_serializer(kpi_photos, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except KPIPhotos.DoesNotExist:
                return Response({"error": "KPI Photo not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
        elif monitoring_id:
            try:
                if 'pk' in kwargs:  # Check if specific instance is requested
                    kpi_photos = KPIPhotos.objects.get(pk=kwargs['pk'])
                    serializer = self.get_serializer(kpi_photos)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else:
                    kpi_photos = self.get_queryset()
                    serializer = self.get_serializer(kpi_photos, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except KPIPhotos.DoesNotExist:
                return Response({"error": "KPI Photo not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        elif is_superadmin:
            try:
                if 'pk' in kwargs:  # Check if specific instance is requested
                    kpi_photos = KPIPhotos.objects.get(pk=kwargs['pk'])
                    serializer = self.get_serializer(kpi_photos)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else:
                    kpi_photos = self.get_queryset()
                    serializer = self.get_serializer(kpi_photos, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except KPIPhotos.DoesNotExist:
                return Response({"error": "KPI Photo not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def post(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        try:
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        

class KPIPhotosRetrieveUpdateDelete(generics.RetrieveUpdateDestroyAPIView):
    queryset = KPIPhotos.objects.all()
    serializer_class = KPIPhotosSerializer
    permission_classes = [IsAuthenticated]

    
    
    def get_kpiphotos(self, pk):
        return get_object_or_404(KPIPhotos, pk=pk)
    
    def delete(self, request, pk):
        try:
            kpiphotos = self.get_kpiphotos(pk)
            kpiphotos.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


    def put(self, request, pk, *args, **kwargs):
        instance = self.get_kpiphotos(pk)
        serializer = self.get_serializer(instance, data=request.data, partial=True)
        try:
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



class ReportSummaryListCreateView(generics.ListCreateAPIView):
    queryset = ReportSummary.objects.all()
    serializer_class = ReportSummarySerializer
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        division_id = request.user.division_id
        sector_id = request.user.sector_id
        monitoring_id = request.user.monitoring_id
        is_superadmin = request.user.is_superadmin
        if division_id:
            try:
                if 'pk' in kwargs: 
                    summary = ReportSummary.objects.get(pk=kwargs['pk'], division_id=division_id)
                    serializer = ReportSummarySerializer(summary)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else: 
                    summaries = ReportSummary.objects.filter(division_id=division_id)
                    serializer = ReportSummarySerializer(summaries, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except ReportSummary.DoesNotExist:
                return Response({"error": "Report Summary not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        elif sector_id:
            try:
                if 'pk' in kwargs: 
                    summary = ReportSummary.objects.get(pk=kwargs['pk'], sector_id=sector_id)
                    serializer = ReportSummarySerializer(summary)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else: 
                    summaries = ReportSummary.objects.filter(sector_id=sector_id)
                    serializer = ReportSummarySerializer(summaries, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except ReportSummary.DoesNotExist:
                return Response({"error": "Report Summary not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
        elif monitoring_id:
            try:
                if 'pk' in kwargs: 
                    summary = ReportSummary.objects.get(pk=kwargs['pk'], monitoring_id=monitoring_id)
                    serializer = ReportSummarySerializer(summary)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else: 
                    summaries = ReportSummary.objects.filter(monitoring_id=monitoring_id)
                    serializer = ReportSummarySerializer(summaries, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except ReportSummary.DoesNotExist:
                return Response({"error": "Report Summary not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
        elif is_superadmin:
            try:
                if 'pk' in kwargs: 
                    summary = ReportSummary.objects.get(pk=kwargs['pk'])
                    serializer = ReportSummarySerializer(summary)
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else: 
                    summaries = ReportSummary.objects.all()
                    serializer = ReportSummarySerializer(summaries, many=True)
                    return Response(serializer.data, status=status.HTTP_200_OK)
            except ReportSummary.DoesNotExist:
                return Response({"error": "Report Summary not found"}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def post(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        try:
            if serializer.is_valid():
                user = request.user
                user_has_required_association = False
                
                if (hasattr(user, 'sector_id') and user.sector_id is not None and hasattr(user, 'division_id') and user.division_id is None):
                    user_has_required_association = True
                    serializer.validated_data['sector_id'] = user.sector_id
                if hasattr(user, 'monitoring_id') and user.monitoring_id is not None:
                    user_has_required_association = True
                    serializer.validated_data['monitoring_id'] = user.monitoring_id
                if hasattr(user, 'division_id') and user.division_id is not None:
                    user_has_required_association = True
                    serializer.validated_data['division_id'] = user.division_id

                if user_has_required_association:
                    serializer.save()
                    return Response(serializer.data, status=status.HTTP_201_CREATED)
                else:
                    return Response({"error": "User must belong to at least one sector_id, monitoring_id, or division_id."}, status=status.HTTP_400_BAD_REQUEST)

            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class ReportSummaryDetailView(generics.RetrieveUpdateDestroyAPIView):
    queryset = ReportSummary.objects.all()
    serializer_class = ReportSummarySerializer
    permission_classes = [IsAuthenticated]

    def get_summary(self, pk):
        return get_object_or_404(ReportSummary, pk=pk)
    
    def delete(self, request, pk):
        try:
            summary = self.get_summary(pk)
            summary.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def put(self, request, pk, *args, **kwargs):
        instance = self.get_summary(pk)
        serializer = self.get_serializer(instance, data=request.data, partial=True)
        try:
            if serializer.is_valid():
                user = request.user
                user_has_required_association = False
                
                if (hasattr(user, 'sector_id') and user.sector_id is not None and hasattr(user, 'division_id') and user.division_id is None):
                    user_has_required_association = True
                    serializer.validated_data['sector_id'] = user.sector_id
                if hasattr(user, 'monitoring_id') and user.monitoring_id is not None:
                    user_has_required_association = True
                    serializer.validated_data['monitoring_id'] = user.monitoring_id
                if hasattr(user, 'division_id') and user.division_id is not None:
                    user_has_required_association = True
                    serializer.validated_data['division_id'] = user.division_id

                if user_has_required_association:
                    serializer.save()
                    return Response(serializer.data, status=status.HTTP_200_OK)
                else:
                    return Response({"error": "User must belong to at least one sector_id, monitoring_id, or division_id."}, status=status.HTTP_400_BAD_REQUEST)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# refactoring the code 

unit_conversion_cache = {}
def get_base_unit_info(measure):
    if measure is None:
        return None, 1.0, ""  # Default values if measure is None
    if measure.id not in unit_conversion_cache:
        try:
            # Find the base unit for the measure
            base_unit = Unit.objects.get(measure_id=measure, isBaseUnit=True)
            unit_conversion_cache[measure.id] = (
                base_unit,
                base_unit.conversionFactor,
                base_unit.symbol or "",
            )
        except Unit.DoesNotExist:
            # Fallback if no base unit is defined (should not happen ideally)
            unit_conversion_cache[measure.id] = (None, 1.0, "")
        except Unit.MultipleObjectsReturned:
            # Handle case where multiple base units are marked for the same measure
            # Log this error and potentially return the first one found
            print(
                f"Warning: Multiple base units found for measure ID {measure.id}. Using the first one."
            )
            base_unit = Unit.objects.filter(measure_id=measure, isBaseUnit=True).first()
            unit_conversion_cache[measure.id] = (
                base_unit,
                base_unit.conversionFactor if base_unit else 1.0,
                base_unit.symbol or "" if base_unit else "",
            )
    return unit_conversion_cache[measure.id]

# Helper function to convert a value to the measure's base unit value
def convert_to_base_value(value, unit_instance):
    if value is None or unit_instance is None or unit_instance.conversionFactor is None:
        return 0.0  # Return 0 if value or unit/factor is missing
    # Ensure conversionFactor is treated as float
    try:
        conversion_factor = float(unit_instance.conversionFactor)
        return float(value) * conversion_factor
    except (ValueError, TypeError):
        print(
            f"Warning: Could not convert value '{value}' or conversion factor '{unit_instance.conversionFactor}' to float for unit {unit_instance.name}"
        )
        return 0.0


# Helper function to calculate performance percentage
def calculate_performance_percent(performance_base_value, plan_base_value):
    if plan_base_value is None or plan_base_value == 0:
        return 0.0  # Avoid division by zero
    if performance_base_value is None:
        performance_base_value = 0.0
    try:
        percent = (float(performance_base_value) / float(plan_base_value)) * 100
        # Cap percentage at 100% if needed, or handle >100% based on requirements
        # return min(percent, 100.0)
        return percent  # Returning actual percentage, can be > 100
    except (ValueError, TypeError):
        print(
            f"Warning: Could not calculate percentage for performance '{performance_base_value}' / plan '{plan_base_value}'"
        )
        return 0.0



# Helper function for styling paragraphs (minor adjustments for clarity)
def set_paragraph_style(
    paragraph,
    font_name="Ebrima",
    font_size=Pt(12),
    line_spacing=1.15,
    space_after=Pt(12),
    space_before=Pt(0),
    bold=False,
    bullet=False,
    arrow_bullet=False,
    color=None,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
):
    """Helper function to set the style of a paragraph."""
    if not paragraph.runs:
        run = paragraph.add_run()
    else:
        # Apply style to all runs if needed, or just the first? Assuming first for now.
        run = paragraph.runs[0]
    run.font.name = font_name
    run.bold = bold
    try:
        run.font.size = font_size
    except ValueError:
        print(f"Warning: Invalid font size {font_size}. Using default.")
        run.font.size = Pt(12)  # Default fallback
    # Ensure font is set for complex scripts (like Amharic)
    r = run._element
    rPr = r.get_or_add_rPr()
    fonts = rPr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rPr.append(fonts)
    fonts.set(qn("w:eastAsia"), font_name)
    fonts.set(qn("w:cs"), font_name)  # Also set for complex scripts

    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = line_spacing
    paragraph_format.space_after = space_after
    paragraph_format.space_before = space_before
    paragraph_format.alignment = alignment
    if bullet:
        paragraph.style = (
            "List Bullet"  # Ensure this style exists in the template or default
        )
    if arrow_bullet:
        # Ensure text exists before prepending
        current_text = paragraph.text
        paragraph.text = f"→ {current_text}"
        # Re-apply style to the modified text if necessary (usually handled by paragraph style)
    if color:
        try:
            run.font.color.rgb = RGBColor(*color) if isinstance(color, tuple) else color
        except Exception as e:
            print(f"Warning: Could not apply color {color}. Error: {e}")




# Helper function to set cell shading
def set_cell_shading(cell, color):
    """Sets background color for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

# Helper function to set cell borders (from original code)
def set_cell_border(cell, **kwargs):
    """Sets borders for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for key, val in kwargs.items():
        if val is None:  # Skip if the border style is None
            continue
        # Ensure the border key is valid (e.g., 'top', 'bottom', 'left', 'right')
        valid_borders = ["top", "bottom", "left", "right", "insideH", "insideV"]
        if key not in valid_borders:
            print(f"Warning: Invalid border key '{key}'. Skipping.")
            continue

        border_tag = f"w:{key}"
        border_elem = tcBorders.find(qn(border_tag))
        if border_elem is None:
            border_elem = OxmlElement(border_tag)
            tcBorders.append(border_elem)

        # Ensure val is a dictionary
        if not isinstance(val, dict):
            print(f"Warning: Border value for '{key}' is not a dict: {val}. Skipping.")
            continue

        for k, v in val.items():
            try:
                border_elem.set(qn(f"w:{k}"), str(v))
            except ValueError as e:
                print(
                    f"Warning: Could not set border attribute w:{k}='{v}' for {key}. Error: {e}"
                )

# Helper function to add pictures safely
def add_picture_to_run(run, image_path, width):
    """Adds a picture to a run, handling potential file errors."""
    try:
        if os.path.exists(image_path):
            run.add_picture(image_path, width=width)
        else:
            print(f"Warning: Image file not found at {image_path}")
    except Exception as e:
        print(f"Error adding picture {image_path}: {e}")



def get_filter_display_name(user, sector_param, division_param, quarter, year):
    """Generate dynamic display names based on filters"""
    quarter_display_names = {
        "first": "የመጀመሪያ ሩብ ዓመት",
        "second": "የሁለተኛ ሩብ ዓመት",
        "third": "የሶስተኛ ሩብ ዓመት",
        "fourth": "የአራተኛ ሩብ ዓመት",
        "six": "የ 6 ወር",
        "nine": "የ 9 ወር",
        "year": "ዓመታዊ",
    }
    
    quarter_name = quarter_display_names.get(quarter or "year", "ዓመታዊ")
    
    org_name = ""
    org_type = ""
    
    if sector_param:
        sectors = Sector.objects.filter(id__in=sector_param.split(','))
        org_name = " እና ".join([s.name for s in sectors])
        org_type = "ሴክተር"
    elif division_param:
        divisions = Division.objects.filter(id__in=division_param.split(','))
        org_name = " እና ".join([d.name for d in divisions])
        org_type = "ዲቪዥን"
    elif getattr(user, 'sector_id', None):
        org_name = user.sector_id.name
        org_type = "ሴክተር"
    elif getattr(user, 'division_id', None):
        org_name = user.division_id.name
        org_type = "ዲቪዥን"
    else:
        org_name = "ኢትዮጵያ አርቴፊሻል ኢንተለጀንስ ኢንስቲትዩት"
        org_type = ""
    
    return org_name, org_type, quarter_name


def build_filters(request):
    """Build comprehensive filters based on user role and query parameters"""
    user = request.user
    year_str = request.query_params.get("year") if hasattr(request, 'query_params') else request.GET.get("year")
    quarter = request.query_params.get("quarter") if hasattr(request, 'query_params') else request.GET.get("quarter")
    sector_param = request.query_params.get("sector") if hasattr(request, 'query_params') else request.GET.get("sector")
    division_param = request.query_params.get("division") if hasattr(request, 'query_params') else request.GET.get("division")
    
    if not year_str or not year_str.isdigit():
        return None, "Valid Year parameter is required"
    year = int(year_str)
    
    valid_quarters = ["first", "second", "third", "fourth", "six", "nine", "year"]
    if quarter and quarter.lower() not in valid_quarters:
        return None, f"Invalid quarter. Choose from: {', '.join(valid_quarters)}"
    quarter = quarter.lower() if quarter else None
    annual_kpi_filter = Q(year=year)
    summary_filter = Q(year=year)
    
    if quarter and quarter != 'year':
        summary_filter &= Q(quarter__iexact=quarter)
    
    if sector_param:
        sector_ids = [int(s) for s in sector_param.split(',')]
        summary_filter &= Q(sector_id__in=sector_ids)
        annual_kpi_filter &= Q(kpi__main_goal_id__sector_id__in=sector_ids)
    elif division_param:
        division_ids = [int(d) for d in division_param.split(',')]
        summary_filter &= Q(division_id__in=division_ids)
        annual_kpi_filter &= Q(division_id__in=division_ids)
    elif getattr(user, 'division_id', None):
        summary_filter &= Q(division_id=user.division_id.id)
        annual_kpi_filter &= Q(division_id=user.division_id.id)
    elif getattr(user, 'sector_id', None):
        summary_filter &= Q(sector_id=user.sector_id.id)
        annual_kpi_filter &= Q(kpi__main_goal_id__sector_id=user.sector_id.id)
    elif not (getattr(user, 'is_superadmin', False) or getattr(user, 'monitoring_id', None)):
        return None, "You do not have permission to view this data."
    
    return {
        'annual_kpi_filter': annual_kpi_filter,
        'summary_filter': summary_filter,
        'year': year,
        'quarter': quarter,
        'sector_param': sector_param,
        'division_param': division_param
    }, None
    

# class GenerateReportDocument(APIView):
#     permission_classes = [IsAuthenticated]

#     def get(self, request):
#         delete_old_documents()
        
#         filters, error = build_filters(request)
#         if error:
#             return JsonResponse({"error": error}, status=400)
        
#         user = request.user
#         year = filters['year']
#         quarter = filters['quarter']
#         annual_kpi_filter = filters['annual_kpi_filter']
#         summary_filter = filters['summary_filter']
        
#         org_name, org_type, quarter_name = get_filter_display_name(
#             user,
#             filters['sector_param'],
#             filters['division_param'],
#             quarter,
#             year
#         )
        

#         annual_kpis_queryset = (
#             AnnualKPI.objects.filter(annual_kpi_filter)
#             .select_related(
#                 "kpi",
#                 "kpi__main_goal_id",
#                 "kpi__main_goal_id__strategic_goal_id",
#                 "measure",
#                 "annual_unit_id",
#                 "division_id",
#                 "initial_unit_id",
#                 "pl1_unit_id", "pl2_unit_id", "pl3_unit_id", "pl4_unit_id",
#                 "pr1_unit_id", "pr2_unit_id", "pr3_unit_id", "pr4_unit_id",
#             )
#             .prefetch_related(
#                 Prefetch(
#                     "kpidescription",
#                     queryset=KPIDescription.objects.prefetch_related(
#                         Prefetch(
#                             "description",
#                             queryset=Description.objects.prefetch_related("description_photo"),
#                         )
#                     ),
#                 )
#             )
#             .order_by(
#                 "kpi__main_goal_id__strategic_goal_id__name",
#                 "kpi__main_goal_id__name",
#                 "kpi__name",
#             )
#         )
        
 
#         summaries_queryset = (
#             Summary.objects.filter(summary_filter)
#             .prefetch_related(
#                 "summary_files",
#                 Prefetch(
#                     "summary_subtitle",
#                     queryset=SummarySubtitle.objects.prefetch_related("summary_photo"),
#                 ),
#             )
#             .order_by("type", "id")
#         )
        
#         if not annual_kpis_queryset.exists() and not summaries_queryset.exists():
#             return JsonResponse(
#                 {"error": f"No data found for Year {year}" + (f" and Quarter {quarter}" if quarter else "")}, 
#                 status=404
#             )
      
#         doc = Document()
        

#         try:
#             system_setting = SystemSetting.objects.first()
#             logo_path = system_setting.logo_image.path if system_setting and system_setting.logo_image else None
#         except Exception as e:
#             print(f"Could not load system settings: {e}")
#             logo_path = None
        
#         section = doc.sections[0]
#         section.top_margin = Inches(1)
#         section.bottom_margin = Inches(1)
        
#         if logo_path and os.path.exists(logo_path):
#             img_p = doc.add_paragraph()
#             add_picture_to_run(img_p.add_run(), logo_path, width=Inches(2))
#             img_p.alignment = 1
        

#         title_text = f"{org_name} {org_type} የ {year} በጀት ዓመት {quarter_name}  አፈፃፀም ሪፖርት"
#         title_p = doc.add_paragraph(title_text)
#         set_paragraph_style(title_p, font_size=Pt(14), bold=True, alignment=1)
        
#         doc.add_page_break()
        
    
#         if summaries_queryset.exists():
#             summary_heading = doc.add_paragraph("የስራ ማጠቃለያ")
#             set_paragraph_style(summary_heading, font_size=Pt(14), bold=True)
            
#             for summary in summaries_queryset:
       
#                 if summary.title:
#                     title_para = doc.add_paragraph(summary.title)
#                     set_paragraph_style(title_para, font_size=Pt(13), bold=True)
                
    
#                 if summary.description:
#                     desc_para = doc.add_paragraph(summary.description)
#                     set_paragraph_style(desc_para, font_size=Pt(12))
                
        
#                 for summary_file in summary.summary_files.all():
#                     if summary_file.photos and os.path.exists(summary_file.photos.path):
#                         img_para = doc.add_paragraph()
#                         add_picture_to_run(img_para.add_run(), summary_file.photos.path, width=Inches(5))
#                         img_para.alignment = 1
                
      
#                 for subtitle in summary.summary_subtitle.all():
#                     if subtitle.subtitle:
#                         sub_para = doc.add_paragraph(subtitle.subtitle)
#                         set_paragraph_style(sub_para, font_size=Pt(12), bold=True)
                    
#                     if subtitle.description:
#                         sub_desc_para = doc.add_paragraph(subtitle.description)
#                         set_paragraph_style(sub_desc_para, font_size=Pt(11))
                    
#                     for photo in subtitle.summary_photo.all():
#                         if photo.photos and os.path.exists(photo.photos.path):
#                             photo_para = doc.add_paragraph()
#                             add_picture_to_run(photo_para.add_run(), photo.photos.path, width=Inches(4))
#                             photo_para.alignment = 1
            
#             doc.add_page_break()
        
#         if annual_kpis_queryset.exists():
#             kpi_desc_heading = doc.add_paragraph("የስራ አፈፃፀም መግለጫዎች")
#             set_paragraph_style(kpi_desc_heading, font_size=Pt(14), bold=True)
            
#             for annual_kpi in annual_kpis_queryset:
#                 kpi_descriptions = annual_kpi.kpidescription.all()
                
#                 if kpi_descriptions.exists():
#                     for kpi_desc in kpi_descriptions:
#                         kpi_heading = doc.add_paragraph(annual_kpi.kpi.name)
#                         set_paragraph_style(kpi_heading, font_size=Pt(13), bold=True)
                        

#                         for desc in kpi_desc.description.all():
#                             if desc.description:
#                                 desc_para = doc.add_paragraph(desc.description)
#                                 set_paragraph_style(desc_para, font_size=Pt(12))
                            
#                             for photo in desc.description_photo.all():
#                                 if photo.photos and os.path.exists(photo.photos.path):
#                                     photo_para = doc.add_paragraph()
#                                     add_picture_to_run(photo_para.add_run(), photo.photos.path, width=Inches(4))
#                                     photo_para.alignment = 1
            
#             doc.add_page_break()
        

#         generate_kpi_performance_table(doc, request, filters)
        
 
#         try:
#             with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
#                 temp_path = tmp_file.name
#                 doc.save(temp_path)
            
#             doc_dir = os.path.join(settings.MEDIA_ROOT, "documents")
#             os.makedirs(doc_dir, exist_ok=True)
            
#             safe_qualifier = re.sub(r"\W+", "_", org_name)
#             base_filename = f"{year}_{quarter or 'annual'}_{safe_qualifier}_report"
#             docx_filename = f"{base_filename}.docx"
#             pdf_filename = f"{base_filename}.pdf"
            
#             docx_rel = os.path.join("documents", docx_filename)
#             pdf_rel = os.path.join("documents", pdf_filename)
#             docx_full = os.path.join(doc_dir, docx_filename)
#             pdf_full = os.path.join(doc_dir, pdf_filename)
            
#             with open(temp_path, "rb") as f:
#                 default_storage.save(docx_rel, f)
#             os.remove(temp_path)
            
#             actual_docx_path = default_storage.path(docx_rel)
#             subprocess.run(
#                 [
#                     "libreoffice",
#                     "--headless",
#                     "--convert-to",
#                     "pdf",
#                     "--outdir",
#                     doc_dir,
#                     actual_docx_path,
#                 ],
#                 check=True,
#                 stdout=subprocess.DEVNULL,
#                 stderr=subprocess.DEVNULL,
#             )
            
#             return JsonResponse(
#                 {
#                     "message": "Document saved successfully",
#                     "docx_file_path": f"http://127.0.0.1:8000/media/documents/{docx_filename}",
#                     "pdf_file_path": f"http://127.0.0.1:8000/media/documents/{pdf_filename}",
#                 },
#                 status=200,
#             )
        
#         except subprocess.CalledProcessError as e:
#             print(f"LibreOffice conversion failed: {e}")
#             return JsonResponse({"error": f"LibreOffice conversion failed: {e}"}, status=500)
#         except Exception as e:
#             print(f"Error during document saving/conversion: {e}")
#             import traceback
#             traceback.print_exc()
#             return JsonResponse({"error": str(e)}, status=500)
class GenerateReportDocumentc(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        delete_old_documents()
        
        filters, error = build_filters(request)
        if error:
            return JsonResponse({"error": error}, status=400)
        
        user = request.user
        year = filters['year']
        quarter = filters['quarter']
        annual_kpi_filter = filters['annual_kpi_filter']
        summary_filter = filters['summary_filter']
        
        org_name, org_type, quarter_name = get_filter_display_name(
            user,
            filters['sector_param'],
            filters['division_param'],
            quarter,
            year
        )
        
        # Apply access control rules to summary filter
        division_id = getattr(user, 'division_id', None)
        sector_id = getattr(user, 'sector_id', None)
        
        # For sector users, they can see summaries from their sector AND its divisions
        if sector_id and not division_id:
            sector_divisions = Division.objects.filter(sector_id=sector_id).values_list('id', flat=True)
            # Update summary_filter to include sector OR its divisions
            summary_filter = summary_filter & (
                Q(sector_id=sector_id) | 
                Q(division_id__in=sector_divisions)
            )
        
        # For division users, they can only see their division
        elif division_id:
            summary_filter = summary_filter & Q(division_id=division_id)
        
        # Monitoring and Superadmin have no restrictions - they can see all
        
        # Apply access control to annual_kpi_filter as well
        if sector_id and not division_id:
            # Sector users can see KPIs from their sector AND its divisions
            sector_divisions = Division.objects.filter(sector_id=sector_id).values_list('id', flat=True)
            annual_kpi_filter = annual_kpi_filter & (
                Q(kpi__main_goal_id__sector_id=sector_id) | 
                Q(division_id__in=sector_divisions)
            )
        elif division_id:
            # Division users can only see their division KPIs
            annual_kpi_filter = annual_kpi_filter & Q(division_id=division_id)

        annual_kpis_queryset = (
            AnnualKPI.objects.filter(annual_kpi_filter)
            .select_related(
                "kpi",
                "kpi__main_goal_id",
                "kpi__main_goal_id__strategic_goal_id",
                "measure",
                "annual_unit_id",
                "division_id",
                "initial_unit_id",
                "pl1_unit_id", "pl2_unit_id", "pl3_unit_id", "pl4_unit_id",
                "pr1_unit_id", "pr2_unit_id", "pr3_unit_id", "pr4_unit_id",
            )
            .prefetch_related(
                Prefetch(
                    "kpidescription",
                    queryset=KPIDescription.objects.prefetch_related(
                        Prefetch(
                            "description",
                            queryset=Description.objects.prefetch_related("description_photo"),
                        )
                    ),
                )
            )
            .order_by(
                "kpi__main_goal_id__strategic_goal_id__name",
                "kpi__main_goal_id__name",
                "kpi__name",
            )
        )
        
        summaries_queryset = (
            Summary.objects.filter(summary_filter)
            .prefetch_related(
                "summary_files",
                Prefetch(
                    "summary_subtitle",
                    queryset=SummarySubtitle.objects.prefetch_related("summary_photo"),
                ),
            )
            .order_by("type", "id")
        )
        
        if not annual_kpis_queryset.exists() and not summaries_queryset.exists():
            return JsonResponse(
                {"error": f"No data found for Year {year}" + (f" and Quarter {quarter}" if quarter else "")}, 
                status=404
            )
      
        doc = Document()
        
        try:
            system_setting = SystemSetting.objects.first()
            logo_path = system_setting.logo_image.path if system_setting and system_setting.logo_image else None
        except Exception as e:
            print(f"Could not load system settings: {e}")
            logo_path = None
        
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        if logo_path and os.path.exists(logo_path):
            img_p = doc.add_paragraph()
            add_picture_to_run(img_p.add_run(), logo_path, width=Inches(2))
            img_p.alignment = 1
        
        title_text = f"{org_name} {org_type} የ {year} በጀት ዓመት {quarter_name}  አፈፃፀም ሪፖርት"
        title_p = doc.add_paragraph(title_text)
        set_paragraph_style(title_p, font_size=Pt(14), bold=True, alignment=1)
        
        doc.add_page_break()
        
        # Define the order of sections
        section_order = [
            ('introduction', 'መግቢያ'),
            ('institutional', 'ተቋማዊ የማስፈጸም አቅም'),
            ('strategic', 'ስትራተጂክ ግብ'),
            ('main_goal', 'ዋና ግብ'),
            ('kpi', 'የአፈፃፀም አመልካች'),
            ('kpi_description', 'የስራ አፈፃፀም መግለጫዎች'),
            ('performance_table', 'የአፈፃፀም ሰንጠረዥ'),
            ('challenges', 'ተግዳሮቶችና የመፍትሔዎች'),
            ('conclusion', 'ማጠቃለያ')
        ]
        
        # Process each section in order
        for section_type, section_title in section_order:
            
            if section_type == 'introduction':
                # መግቢያ section
                intro_summaries = summaries_queryset.filter(type='introduction')
                if intro_summaries.exists():
                    intro_heading = doc.add_paragraph("መግቢያ")
                    set_paragraph_style(intro_heading, font_size=Pt(14), bold=True)
                    
                    for summary in intro_summaries:
                        if summary.title:
                            title_para = doc.add_paragraph(summary.title)
                            set_paragraph_style(title_para, font_size=Pt(13), bold=True)
                        
                        if summary.description:
                            desc_para = doc.add_paragraph(summary.description)
                            set_paragraph_style(desc_para, font_size=Pt(12))
                        
                        for summary_file in summary.summary_files.all():
                            if summary_file.photos and os.path.exists(summary_file.photos.path):
                                img_para = doc.add_paragraph()
                                add_picture_to_run(img_para.add_run(), summary_file.photos.path, width=Inches(5))
                                img_para.alignment = 1
                        
                        for subtitle in summary.summary_subtitle.all():
                            if subtitle.subtitle:
                                sub_para = doc.add_paragraph(subtitle.subtitle)
                                set_paragraph_style(sub_para, font_size=Pt(12), bold=True)
                            
                            if subtitle.description:
                                sub_desc_para = doc.add_paragraph(subtitle.description)
                                set_paragraph_style(sub_desc_para, font_size=Pt(11))
                            
                            for photo in subtitle.summary_photo.all():
                                if photo.photos and os.path.exists(photo.photos.path):
                                    photo_para = doc.add_paragraph()
                                    add_picture_to_run(photo_para.add_run(), photo.photos.path, width=Inches(4))
                                    photo_para.alignment = 1
                    
                    doc.add_page_break()
            
            elif section_type == 'institutional':
                # ተቋማዊ የማስፈጸም አቅም section
                institutional_summaries = summaries_queryset.filter(type='institutional')
                if institutional_summaries.exists():
                    inst_heading = doc.add_paragraph("ተቋማዊ የማስፈጸም አቅም")
                    set_paragraph_style(inst_heading, font_size=Pt(14), bold=True)
                    
                    for summary in institutional_summaries:
                        if summary.title:
                            title_para = doc.add_paragraph(summary.title)
                            set_paragraph_style(title_para, font_size=Pt(13), bold=True)
                        
                        if summary.description:
                            desc_para = doc.add_paragraph(summary.description)
                            set_paragraph_style(desc_para, font_size=Pt(12))
                        
                        for summary_file in summary.summary_files.all():
                            if summary_file.photos and os.path.exists(summary_file.photos.path):
                                img_para = doc.add_paragraph()
                                add_picture_to_run(img_para.add_run(), summary_file.photos.path, width=Inches(5))
                                img_para.alignment = 1
                        
                        for subtitle in summary.summary_subtitle.all():
                            if subtitle.subtitle:
                                sub_para = doc.add_paragraph(subtitle.subtitle)
                                set_paragraph_style(sub_para, font_size=Pt(12), bold=True)
                            
                            if subtitle.description:
                                sub_desc_para = doc.add_paragraph(subtitle.description)
                                set_paragraph_style(sub_desc_para, font_size=Pt(11))
                            
                            for photo in subtitle.summary_photo.all():
                                if photo.photos and os.path.exists(photo.photos.path):
                                    photo_para = doc.add_paragraph()
                                    add_picture_to_run(photo_para.add_run(), photo.photos.path, width=Inches(4))
                                    photo_para.alignment = 1
                    
                    doc.add_page_break()
            
            elif section_type == 'strategic':
                # Group by strategic goal for structured display
                strategic_goals = {}
                for annual_kpi in annual_kpis_queryset:
                    strategic_goal = annual_kpi.kpi.main_goal_id.strategic_goal_id
                    if strategic_goal not in strategic_goals:
                        strategic_goals[strategic_goal] = []
                    strategic_goals[strategic_goal].append(annual_kpi)
                
                if strategic_goals:
                    strategic_heading = doc.add_paragraph("ስትራተጂክ ግብ")
                    set_paragraph_style(strategic_heading, font_size=Pt(14), bold=True)
                    
                    for i, (strategic_goal, kpis) in enumerate(strategic_goals.items(), 1):
                        goal_para = doc.add_paragraph(f"{i}. {strategic_goal.name}")
                        set_paragraph_style(goal_para, font_size=Pt(13), bold=True)
                    
                    doc.add_page_break()
            
            elif section_type == 'main_goal':
                # Group by main goal under each strategic goal
                from collections import defaultdict
                strategic_main_goals = defaultdict(lambda: defaultdict(list))
                
                for annual_kpi in annual_kpis_queryset:
                    strategic_goal = annual_kpi.kpi.main_goal_id.strategic_goal_id
                    main_goal = annual_kpi.kpi.main_goal_id
                    strategic_main_goals[strategic_goal][main_goal].append(annual_kpi)
                
                if strategic_main_goals:
                    main_goal_heading = doc.add_paragraph("ዋና ግብ")
                    set_paragraph_style(main_goal_heading, font_size=Pt(14), bold=True)
                    
                    strategic_counter = 0
                    for strategic_goal, main_goals in strategic_main_goals.items():
                        strategic_counter += 1
                        for j, (main_goal, kpis) in enumerate(main_goals.items(), 1):
                            main_goal_para = doc.add_paragraph(f"{strategic_counter}.{j}. {main_goal.name}")
                            set_paragraph_style(main_goal_para, font_size=Pt(12), bold=True)
                    
                    doc.add_page_break()
            
            elif section_type == 'kpi':
                # List KPIs under each main goal
                from collections import defaultdict
                main_goal_kpis = defaultdict(list)
                
                for annual_kpi in annual_kpis_queryset:
                    main_goal = annual_kpi.kpi.main_goal_id
                    main_goal_kpis[main_goal].append(annual_kpi)
                
                if main_goal_kpis:
                    kpi_heading = doc.add_paragraph("የአፈፃፀም አመልካች")
                    set_paragraph_style(kpi_heading, font_size=Pt(14), bold=True)
                    
                    # Get strategic goal mapping for numbering
                    strategic_goals = {}
                    strategic_counter = 0
                    for annual_kpi in annual_kpis_queryset:
                        strategic_goal = annual_kpi.kpi.main_goal_id.strategic_goal_id
                        if strategic_goal not in strategic_goals:
                            strategic_counter += 1
                            strategic_goals[strategic_goal] = strategic_counter
                    
                    # Reset counters for structured numbering
                    main_goal_counter_map = {}
                    kpi_counter_map = {}
                    
                    for annual_kpi in annual_kpis_queryset:
                        strategic_goal = annual_kpi.kpi.main_goal_id.strategic_goal_id
                        main_goal = annual_kpi.kpi.main_goal_id
                        
                        strategic_num = strategic_goals[strategic_goal]
                        
                        if main_goal not in main_goal_counter_map:
                            main_goal_counter_map[main_goal] = len(main_goal_counter_map) + 1
                        main_num = main_goal_counter_map[main_goal]
                        
                        if annual_kpi.kpi not in kpi_counter_map:
                            kpi_counter_map[annual_kpi.kpi] = len(kpi_counter_map) + 1
                        kpi_num = kpi_counter_map[annual_kpi.kpi]
                        
                        kpi_para = doc.add_paragraph(f"{strategic_num}.{main_num}.{kpi_num}. {annual_kpi.kpi.name}")
                        set_paragraph_style(kpi_para, font_size=Pt(12))
                    
                    doc.add_page_break()
            
            elif section_type == 'kpi_description':
                # KPI descriptions section
                if annual_kpis_queryset.exists():
                    kpi_desc_heading = doc.add_paragraph("የስራ አፈፃፀም መግለጫዎች")
                    set_paragraph_style(kpi_desc_heading, font_size=Pt(14), bold=True)
                    
                    # Group by KPI for description display
                    kpi_descriptions_map = {}
                    for annual_kpi in annual_kpis_queryset:
                        kpi_descriptions = annual_kpi.kpidescription.all()
                        if kpi_descriptions.exists():
                            kpi_descriptions_map[annual_kpi.kpi] = kpi_descriptions
                    
                    # Get structured numbering
                    strategic_goals = {}
                    strategic_counter = 0
                    for annual_kpi in annual_kpis_queryset:
                        strategic_goal = annual_kpi.kpi.main_goal_id.strategic_goal_id
                        if strategic_goal not in strategic_goals:
                            strategic_counter += 1
                            strategic_goals[strategic_goal] = strategic_counter
                    
                    main_goal_counter_map = {}
                    kpi_counter_map = {}
                    
                    for kpi, descriptions in kpi_descriptions_map.items():
                        strategic_goal = kpi.main_goal_id.strategic_goal_id
                        main_goal = kpi.main_goal_id
                        
                        strategic_num = strategic_goals[strategic_goal]
                        
                        if main_goal not in main_goal_counter_map:
                            main_goal_counter_map[main_goal] = len(main_goal_counter_map) + 1
                        main_num = main_goal_counter_map[main_goal]
                        
                        if kpi not in kpi_counter_map:
                            kpi_counter_map[kpi] = len(kpi_counter_map) + 1
                        kpi_num = kpi_counter_map[kpi]
                        
                        kpi_heading = doc.add_paragraph(f"{strategic_num}.{main_num}.{kpi_num}. {kpi.name}")
                        set_paragraph_style(kpi_heading, font_size=Pt(13), bold=True)
                        
                        for kpi_desc in descriptions:
                            for desc in kpi_desc.description.all():
                                if desc.description:
                                    desc_para = doc.add_paragraph(desc.description)
                                    set_paragraph_style(desc_para, font_size=Pt(12))
                                
                                for photo in desc.description_photo.all():
                                    if photo.photos and os.path.exists(photo.photos.path):
                                        photo_para = doc.add_paragraph()
                                        add_picture_to_run(photo_para.add_run(), photo.photos.path, width=Inches(4))
                                        photo_para.alignment = 1
                    
                    doc.add_page_break()
            
            elif section_type == 'performance_table':
                # Performance table section
                if annual_kpis_queryset.exists():
                    table_heading = doc.add_paragraph("የአፈፃፀም ሰንጠረዥ")
                    set_paragraph_style(table_heading, font_size=Pt(14), bold=True)
                    
                    generate_kpi_performance_table(doc, request, filters)
                    
                    doc.add_page_break()
            
            elif section_type == 'challenges':
                # ተግዳሮቶችና የመፍትሔዎች section
                challenges_summaries = summaries_queryset.filter(type='challenges')
                if challenges_summaries.exists():
                    challenges_heading = doc.add_paragraph("ተግዳሮቶችና የመፍትሔዎች")
                    set_paragraph_style(challenges_heading, font_size=Pt(14), bold=True)
                    
                    for summary in challenges_summaries:
                        if summary.title:
                            title_para = doc.add_paragraph(summary.title)
                            set_paragraph_style(title_para, font_size=Pt(13), bold=True)
                        
                        if summary.description:
                            desc_para = doc.add_paragraph(summary.description)
                            set_paragraph_style(desc_para, font_size=Pt(12))
                        
                        for summary_file in summary.summary_files.all():
                            if summary_file.photos and os.path.exists(summary_file.photos.path):
                                img_para = doc.add_paragraph()
                                add_picture_to_run(img_para.add_run(), summary_file.photos.path, width=Inches(5))
                                img_para.alignment = 1
                        
                        for subtitle in summary.summary_subtitle.all():
                            if subtitle.subtitle:
                                sub_para = doc.add_paragraph(subtitle.subtitle)
                                set_paragraph_style(sub_para, font_size=Pt(12), bold=True)
                            
                            if subtitle.description:
                                sub_desc_para = doc.add_paragraph(subtitle.description)
                                set_paragraph_style(sub_desc_para, font_size=Pt(11))
                            
                            for photo in subtitle.summary_photo.all():
                                if photo.photos and os.path.exists(photo.photos.path):
                                    photo_para = doc.add_paragraph()
                                    add_picture_to_run(photo_para.add_run(), photo.photos.path, width=Inches(4))
                                    photo_para.alignment = 1
                    
                    doc.add_page_break()
            
            elif section_type == 'conclusion':
                # ማጠቃለያ section
                conclusion_summaries = summaries_queryset.filter(type='conclusion')
                if conclusion_summaries.exists():
                    conclusion_heading = doc.add_paragraph("ማጠቃለያ")
                    set_paragraph_style(conclusion_heading, font_size=Pt(14), bold=True)
                    
                    for summary in conclusion_summaries:
                        if summary.title:
                            title_para = doc.add_paragraph(summary.title)
                            set_paragraph_style(title_para, font_size=Pt(13), bold=True)
                        
                        if summary.description:
                            desc_para = doc.add_paragraph(summary.description)
                            set_paragraph_style(desc_para, font_size=Pt(12))
                        
                        for summary_file in summary.summary_files.all():
                            if summary_file.photos and os.path.exists(summary_file.photos.path):
                                img_para = doc.add_paragraph()
                                add_picture_to_run(img_para.add_run(), summary_file.photos.path, width=Inches(5))
                                img_para.alignment = 1
                        
                        for subtitle in summary.summary_subtitle.all():
                            if subtitle.subtitle:
                                sub_para = doc.add_paragraph(subtitle.subtitle)
                                set_paragraph_style(sub_para, font_size=Pt(12), bold=True)
                            
                            if subtitle.description:
                                sub_desc_para = doc.add_paragraph(subtitle.description)
                                set_paragraph_style(sub_desc_para, font_size=Pt(11))
                            
                            for photo in subtitle.summary_photo.all():
                                if photo.photos and os.path.exists(photo.photos.path):
                                    photo_para = doc.add_paragraph()
                                    add_picture_to_run(photo_para.add_run(), photo.photos.path, width=Inches(4))
                                    photo_para.alignment = 1
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                temp_path = tmp_file.name
                doc.save(temp_path)
            
            doc_dir = os.path.join(settings.MEDIA_ROOT, "documents")
            os.makedirs(doc_dir, exist_ok=True)
            
            safe_qualifier = re.sub(r"\W+", "_", org_name)
            base_filename = f"{year}_{quarter or 'annual'}_{safe_qualifier}_report"
            docx_filename = f"{base_filename}.docx"
            pdf_filename = f"{base_filename}.pdf"
            
            docx_rel = os.path.join("documents", docx_filename)
            pdf_rel = os.path.join("documents", pdf_filename)
            docx_full = os.path.join(doc_dir, docx_filename)
            pdf_full = os.path.join(doc_dir, pdf_filename)
            
            with open(temp_path, "rb") as f:
                default_storage.save(docx_rel, f)
            os.remove(temp_path)
            
            actual_docx_path = default_storage.path(docx_rel)
            subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    doc_dir,
                    actual_docx_path,
                ],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            
            return JsonResponse(
                {
                    "message": "Document saved successfully",
                    "docx_file_path": f"http://196.188.240.102:4020/media/documents/{docx_filename}",
                    "pdf_file_path": f"http://196.188.240.102:4020/media/documents/{pdf_filename}",
                },
                status=200,
            )
        
        except subprocess.CalledProcessError as e:
            print(f"LibreOffice conversion failed: {e}")
            return JsonResponse({"error": f"LibreOffice conversion failed: {e}"}, status=500)
        except Exception as e:
            print(f"Error during document saving/conversion: {e}")
            import traceback
            traceback.print_exc()
            return JsonResponse({"error": str(e)}, status=500)

class GenerateReportDocument(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        delete_old_documents()
        
        filters, error = build_filters(request)
        if error:
            return JsonResponse({"error": error}, status=400)
        
        user = request.user
        year = filters['year']
        quarter = filters['quarter']
        annual_kpi_filter = filters['annual_kpi_filter']
        summary_filter = filters['summary_filter']
        
        org_name, org_type, quarter_name = get_filter_display_name(
            user,
            filters['sector_param'],
            filters['division_param'],
            quarter,
            year
        )
        
        division_id = getattr(user, 'division_id', None)
        sector_id = getattr(request.user, 'sector_id', None)
        monitoring_id = getattr(request.user, 'monitoring_id', None)
        is_superadmin = getattr(request.user, 'is_superadmin', False)
        
        if sector_id and not division_id:
            from userApp.models import Division
            from django.db.models import Q
            sector_divisions = Division.objects.filter(sector_id=sector_id).values_list('id', flat=True)
            summary_filter = summary_filter & (
                Q(sector_id=sector_id) | 
                Q(division_id__in=sector_divisions)
            )
        

        elif division_id:
            from django.db.models import Q
            summary_filter = summary_filter & Q(division_id=division_id)
        
        if sector_id and not division_id:
            from userApp.models import Division
            from django.db.models import Q
            sector_divisions = Division.objects.filter(sector_id=sector_id).values_list('id', flat=True)
            annual_kpi_filter = annual_kpi_filter & (
                Q(kpi__main_goal_id__sector_id=sector_id) | 
                Q(division_id__in=sector_divisions)
            )
        elif division_id:
            from django.db.models import Q
            annual_kpi_filter = annual_kpi_filter & Q(division_id=division_id)

        annual_kpis_queryset = (
            AnnualKPI.objects.filter(annual_kpi_filter)
            .select_related(
                "kpi",
                "kpi__main_goal_id",
                "kpi__main_goal_id__strategic_goal_id",
                "measure",
                "annual_unit_id",
                "division_id",
                "initial_unit_id",
                "pl1_unit_id", "pl2_unit_id", "pl3_unit_id", "pl4_unit_id",
                "pr1_unit_id", "pr2_unit_id", "pr3_unit_id", "pr4_unit_id",
            )
            .prefetch_related(
                Prefetch(
                    "kpidescription",
                    queryset=KPIDescription.objects.prefetch_related(
                        Prefetch(
                            "description",
                            queryset=Description.objects.prefetch_related("description_photo"),
                        )
                    ),
                )
            )
            .order_by(
                "kpi__main_goal_id__strategic_goal_id__name",
                "kpi__main_goal_id__name",
                "kpi__name",
            )
        )
        
        summaries_queryset = (
            Summary.objects.filter(summary_filter)
            .prefetch_related(
                "summary_files",
                Prefetch(
                    "summary_subtitle",
                    queryset=SummarySubtitle.objects.prefetch_related("summary_photo"),
                ),
            )
            .order_by("id")
        )
        
        if not annual_kpis_queryset.exists() and not summaries_queryset.exists():
            return JsonResponse(
                {"error": f"No data found for Year {year}" + (f" and Quarter {quarter}" if quarter else "")}, 
                status=404
            )
      
        doc = Document()
        
        try:
            system_setting = SystemSetting.objects.first()
            logo_path = system_setting.logo_image.path if system_setting and system_setting.logo_image else None
        except Exception as e:
            print(f"Could not load system settings: {e}")
            logo_path = None
        
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        if logo_path and os.path.exists(logo_path):
            img_p = doc.add_paragraph()
            add_picture_to_run(img_p.add_run(), logo_path, width=Inches(2))
            img_p.alignment = 1
        
        title_text = f"{org_name} {org_type} የ {year} በጀት ዓመት {quarter_name}  አፈፃፀም ሪፖርት"
        title_p = doc.add_paragraph(title_text)
        set_paragraph_style(title_p, font_size=Pt(14), bold=True, alignment=1)
        
        doc.add_page_break()
        
        section_titles = [
            ('introduction', 'መግቢያ'),
            ('institutional', 'ተቋማዊ የማስፈጸም አቅም'),
            ('strategic_goals', 'ስትራተጂክ ግብ እና ተዛማጅ የስራ አፈፃፀም'),
            ('performance_table', 'የአፈፃፀም ሰንጠረዥ'),
            ('challenges', 'ተግዳሮቶችና የመፍትሔዎች'),
            ('conclusion', 'ማጠቃለያ')
        ]
    
        for section_key, section_title in section_titles:
            
            if section_key == 'introduction':
                intro_summaries = summaries_queryset.filter(title='መግቢያ')
                if intro_summaries.exists():
                    intro_heading = doc.add_paragraph("መግቢያ")
                    set_paragraph_style(intro_heading, font_size=Pt(14), bold=True)
                    
                    for summary in intro_summaries:
                        if summary.description:
                            desc_para = doc.add_paragraph(summary.description)
                            set_paragraph_style(desc_para, font_size=Pt(12))
                        
                        for summary_file in summary.summary_files.all():
                            if summary_file.photos and os.path.exists(summary_file.photos.path):
                                img_para = doc.add_paragraph()
                                add_picture_to_run(img_para.add_run(), summary_file.photos.path, width=Inches(5))
                                img_para.alignment = 1
                        
                        for subtitle in summary.summary_subtitle.all():
                            if subtitle.subtitle:
                                sub_para = doc.add_paragraph(subtitle.subtitle)
                                set_paragraph_style(sub_para, font_size=Pt(12), bold=True)
                            
                            if subtitle.description:
                                sub_desc_para = doc.add_paragraph(subtitle.description)
                                set_paragraph_style(sub_desc_para, font_size=Pt(11))
                            
                            for photo in subtitle.summary_photo.all():
                                if photo.photos and os.path.exists(photo.photos.path):
                                    photo_para = doc.add_paragraph()
                                    add_picture_to_run(photo_para.add_run(), photo.photos.path, width=Inches(4))
                                    photo_para.alignment = 1
                    
                    doc.add_page_break()
            
            elif section_key == 'institutional':
                institutional_summaries = summaries_queryset.filter(title='ተቋማዊ የማስፈጸም አቅም')
                if institutional_summaries.exists():
                    inst_heading = doc.add_paragraph("ተቋማዊ የማስፈጸም አቅም")
                    set_paragraph_style(inst_heading, font_size=Pt(14), bold=True)
                    
                    for summary in institutional_summaries:
                        if summary.description:
                            desc_para = doc.add_paragraph(summary.description)
                            set_paragraph_style(desc_para, font_size=Pt(12))
                        
                        for summary_file in summary.summary_files.all():
                            if summary_file.photos and os.path.exists(summary_file.photos.path):
                                img_para = doc.add_paragraph()
                                add_picture_to_run(img_para.add_run(), summary_file.photos.path, width=Inches(5))
                                img_para.alignment = 1
                        
                        for subtitle in summary.summary_subtitle.all():
                            if subtitle.subtitle:
                                sub_para = doc.add_paragraph(subtitle.subtitle)
                                set_paragraph_style(sub_para, font_size=Pt(12), bold=True)
                            
                            if subtitle.description:
                                sub_desc_para = doc.add_paragraph(subtitle.description)
                                set_paragraph_style(sub_desc_para, font_size=Pt(11))
                            
                            for photo in subtitle.summary_photo.all():
                                if photo.photos and os.path.exists(photo.photos.path):
                                    photo_para = doc.add_paragraph()
                                    add_picture_to_run(photo_para.add_run(), photo.photos.path, width=Inches(4))
                                    photo_para.alignment = 1
                    
                    doc.add_page_break()
            
            elif section_key == 'strategic_goals':
                from collections import defaultdict
                strategic_goals_dict = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))
                
                for annual_kpi in annual_kpis_queryset:
                    strategic_goal = annual_kpi.kpi.main_goal_id.strategic_goal_id
                    main_goal = annual_kpi.kpi.main_goal_id
                    kpi = annual_kpi.kpi
                    kpi_descriptions = list(annual_kpi.kpidescription.all())
                    strategic_goals_dict[strategic_goal][main_goal][kpi].extend(kpi_descriptions)
                
                if strategic_goals_dict:
                    strategic_heading = doc.add_paragraph("ስትራተጂክ ግብ እና ተዛማጅ የስራ አፈፃፀም")
                    set_paragraph_style(strategic_heading, font_size=Pt(14), bold=True)
                    strategic_counter = 0
                    for strategic_goal, main_goals_dict in strategic_goals_dict.items():
                        strategic_counter += 1
                        strategic_goal_para = doc.add_paragraph(f"{strategic_counter}. {strategic_goal.name}")
                        set_paragraph_style(strategic_goal_para, font_size=Pt(13), bold=True)
            
                        main_goal_counter = 0
                        for main_goal, kpis_dict in main_goals_dict.items():
                            main_goal_counter += 1
                            
        
                            main_goal_para = doc.add_paragraph(f"   {strategic_counter}.{main_goal_counter}. {main_goal.name}")
                            set_paragraph_style(main_goal_para, font_size=Pt(12), bold=True)
                            
    
                            kpi_counter = 0
                            for kpi, kpi_descriptions in kpis_dict.items():
                                kpi_counter += 1
                                
     
                                kpi_para = doc.add_paragraph(f"      {strategic_counter}.{main_goal_counter}.{kpi_counter}. {kpi.name}")
                                set_paragraph_style(kpi_para, font_size=Pt(11), bold=False)

                                for kpi_desc in kpi_descriptions:
                                    for desc in kpi_desc.description.all():
                                        if desc.description:
                                            desc_text = "         " + desc.description
                                            desc_para = doc.add_paragraph(desc_text)
                                            set_paragraph_style(desc_para, font_size=Pt(11))
                                        
                                        for photo in desc.description_photo.all():
                                            if photo.photos and os.path.exists(photo.photos.path):
                                                photo_para = doc.add_paragraph()
                                                photo_run = photo_para.add_run("         ")
                                                photo_run.font.size = Pt(11)
                                                add_picture_to_run(photo_para.add_run(), photo.photos.path, width=Inches(4))
                                                photo_para.alignment = 1
                    
                    doc.add_page_break()
            
            elif section_key == 'performance_table':
                if annual_kpis_queryset.exists():
                    table_heading = doc.add_paragraph("የአፈፃፀም ሰንጠረዥ")
                    set_paragraph_style(table_heading, font_size=Pt(14), bold=True)
                    
                    generate_kpi_performance_table(doc, request, filters)
                    
                    doc.add_page_break()
            
            elif section_key == 'challenges':
                challenges_summaries = summaries_queryset.filter(title='ተግዳሮቶችና የመፍትሔዎች')
                if challenges_summaries.exists():
                    challenges_heading = doc.add_paragraph("ተግዳሮቶችና የመፍትሔዎች")
                    set_paragraph_style(challenges_heading, font_size=Pt(14), bold=True)
                    
                    for summary in challenges_summaries:
                        if summary.description:
                            desc_para = doc.add_paragraph(summary.description)
                            set_paragraph_style(desc_para, font_size=Pt(12))
                        
                        for summary_file in summary.summary_files.all():
                            if summary_file.photos and os.path.exists(summary_file.photos.path):
                                img_para = doc.add_paragraph()
                                add_picture_to_run(img_para.add_run(), summary_file.photos.path, width=Inches(5))
                                img_para.alignment = 1
                        
                        for subtitle in summary.summary_subtitle.all():
                            if subtitle.subtitle:
                                sub_para = doc.add_paragraph(subtitle.subtitle)
                                set_paragraph_style(sub_para, font_size=Pt(12), bold=True)
                            
                            if subtitle.description:
                                sub_desc_para = doc.add_paragraph(subtitle.description)
                                set_paragraph_style(sub_desc_para, font_size=Pt(11))
                            
                            for photo in subtitle.summary_photo.all():
                                if photo.photos and os.path.exists(photo.photos.path):
                                    photo_para = doc.add_paragraph()
                                    add_picture_to_run(photo_para.add_run(), photo.photos.path, width=Inches(4))
                                    photo_para.alignment = 1
                    
                    doc.add_page_break()
            
            elif section_key == 'conclusion':
                conclusion_summaries = summaries_queryset.filter(title='ማጠቃለያ')
                if conclusion_summaries.exists():
                    conclusion_heading = doc.add_paragraph("ማጠቃለያ")
                    set_paragraph_style(conclusion_heading, font_size=Pt(14), bold=True)
                    
                    for summary in conclusion_summaries:
                        if summary.description:
                            desc_para = doc.add_paragraph(summary.description)
                            set_paragraph_style(desc_para, font_size=Pt(12))
                        
                        for summary_file in summary.summary_files.all():
                            if summary_file.photos and os.path.exists(summary_file.photos.path):
                                img_para = doc.add_paragraph()
                                add_picture_to_run(img_para.add_run(), summary_file.photos.path, width=Inches(5))
                                img_para.alignment = 1
                        
                        for subtitle in summary.summary_subtitle.all():
                            if subtitle.subtitle:
                                sub_para = doc.add_paragraph(subtitle.subtitle)
                                set_paragraph_style(sub_para, font_size=Pt(12), bold=True)
                            
                            if subtitle.description:
                                sub_desc_para = doc.add_paragraph(subtitle.description)
                                set_paragraph_style(sub_desc_para, font_size=Pt(11))
                            
                            for photo in subtitle.summary_photo.all():
                                if photo.photos and os.path.exists(photo.photos.path):
                                    photo_para = doc.add_paragraph()
                                    add_picture_to_run(photo_para.add_run(), photo.photos.path, width=Inches(4))
                                    photo_para.alignment = 1
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                temp_path = tmp_file.name
                doc.save(temp_path)
            
            doc_dir = os.path.join(settings.MEDIA_ROOT, "documents")
            os.makedirs(doc_dir, exist_ok=True)
            
            safe_qualifier = re.sub(r"\W+", "_", org_name)
            base_filename = f"{year}_{quarter or 'annual'}_{safe_qualifier}_report"
            docx_filename = f"{base_filename}.docx"
            pdf_filename = f"{base_filename}.pdf"
            
            docx_rel = os.path.join("documents", docx_filename)
            pdf_rel = os.path.join("documents", pdf_filename)
            docx_full = os.path.join(doc_dir, docx_filename)
            pdf_full = os.path.join(doc_dir, pdf_filename)
            
            with open(temp_path, "rb") as f:
                default_storage.save(docx_rel, f)
            os.remove(temp_path)
            
            actual_docx_path = default_storage.path(docx_rel)
            subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    doc_dir,
                    actual_docx_path,
                ],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            
            return JsonResponse(
                {
                    "message": "Document saved successfully",
                    "docx_file_path": f"http://196.188.240.102:4020/media/documents/{docx_filename}",
                    "pdf_file_path": f"http://196.188.240.102:4020/media/documents/{pdf_filename}",
                },
                status=200,
            )
        
        except subprocess.CalledProcessError as e:
            print(f"LibreOffice conversion failed: {e}")
            return JsonResponse({"error": f"LibreOffice conversion failed: {e}"}, status=500)
        except Exception as e:
            print(f"Error during document saving/conversion: {e}")
            import traceback
            traceback.print_exc()
            return JsonResponse({"error": str(e)}, status=500)


def generate_kpi_performance_table(doc, request, filters=None):
    """Generate KPI performance table with correct indexing"""
    

    if filters is None:
        filters, error = build_filters(request)
        if error:
            return
    
    user = request.user
    year = filters['year']
    quarter = filters['quarter']
    annual_kpi_filter = filters['annual_kpi_filter']
    

    org_name, org_type, quarter_name = get_filter_display_name(
        user,
        filters['sector_param'],
        filters['division_param'],
        quarter,
        year
    )
    
 
    filtered_kpis = AnnualKPI.objects.filter(annual_kpi_filter).select_related(
        'kpi', 'kpi__main_goal_id', 'kpi__main_goal_id__strategic_goal_id',
        'measure', 'division_id'
    )
    
    if not filtered_kpis.exists():
        return
    
    merged_kpis = merge_annual_kpis(filtered_kpis)
    
  
    title_text = f"የ {year} በጀት ዓመት {quarter_name} የ {org_name} {org_type} የስራ አፈፃፀም"
    p = doc.add_paragraph(title_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
 
    table = doc.add_table(rows=2, cols=7)
    table.style = 'Table Grid'
    table.autofit = False
    
   
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ተ.ቁ'
    hdr_cells[1].text = 'ስትራተጂክ ፍጣች፣ ስራ  ተግባራት እና ውጤት  የአፈፃፀም አመልካቶች'
    hdr_cells[2].text = 'መለኪያ'
    hdr_cells[3].text = 'እስከ ' + str(int(year) - 1) + ' አፈፃፀም መንሻ'
    
    merged_report = hdr_cells[4].merge(hdr_cells[5]).merge(hdr_cells[6])
    merged_report.text = f"የ {year} በጀት ዓመት {quarter_name} የስራ  አፈፃፀም"
    
    sub_hdr_cells = table.rows[1].cells
    sub_hdr_cells[4].text = 'ፈፅብ'
    sub_hdr_cells[5].text = 'አፈፃፀም'
    sub_hdr_cells[6].text = 'አፈፃፀም በመቶኛ'
    
    hdr_cells[0].merge(sub_hdr_cells[0])
    hdr_cells[1].merge(sub_hdr_cells[1])
    hdr_cells[2].merge(sub_hdr_cells[2])
    hdr_cells[3].merge(sub_hdr_cells[3])
    
    for row in [table.rows[0], table.rows[1]]:
        for cell in row.cells:
            cell._tc.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="110d7c"/>'.format(nsdecls('w')))
            )
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Ebrima'
                    run.font.color.rgb = RGBColor(255, 188, 20)
    

    strategic_goals = {}
    for kpi in merged_kpis:
        strategic_goal = kpi.kpi.main_goal_id.strategic_goal_id
        main_activity = kpi.kpi.main_goal_id
        
        if strategic_goal not in strategic_goals:
            strategic_goals[strategic_goal] = {}
        if main_activity not in strategic_goals[strategic_goal]:
            strategic_goals[strategic_goal][main_activity] = []
        
        strategic_goals[strategic_goal][main_activity].append(kpi)
    

    strategic_number = 0
    for strategic_goal, main_activities in strategic_goals.items():
        strategic_number += 1
        
  
        row = table.add_row().cells
        row[0].text = f"{strategic_number}"
        row[1].merge(row[6])
        row[1].text = strategic_goal.name
        
   
        set_cell_bg_color(row[0], 'ffbc14')
        set_cell_bg_color(row[1], 'ffbc14')
        set_cell_text_color(row[0], RGBColor(17, 13, 124))
        set_cell_text_color(row[1], RGBColor(17, 13, 124))
        
        main_activity_number = 0
        for main_activity, kpis in main_activities.items():
            main_activity_number += 1
            row = table.add_row().cells
            row[0].text = f"{strategic_number}.{main_activity_number}"
            row[1].merge(row[6])
            row[1].text = main_activity.name
            
            set_cell_bg_color(row[0], 'D3D3D3')
            set_cell_bg_color(row[1], 'D3D3D3')
            set_cell_text_style(row[0])
            set_cell_text_style(row[1])
            
            kpi_number = 0
            for kp in kpis:
                kpi_number += 1
                
                quarter_fields = get_quarter_fields(quarter)
                total_plan, total_perf, symbol, init_base_unit = calculate_kpi_totals(
                    kp, quarter_fields
                )
                
                perf_percent = (total_perf / total_plan * 100) if total_plan > 0 else 0
                perf_percent = round(perf_percent, 2)
                
                row = table.add_row().cells
                row[0].text = f"{strategic_number}.{main_activity_number}.{kpi_number}"
                row[1].text = kp.kpi.name
                row[2].text = f"{kp.measure}"
                row[3].text = f"{round(init_base_unit, 2)}{symbol}"
                row[4].text = f"{round(total_plan, 2)}{symbol}"
                row[5].text = f"{round(total_perf, 2)}{symbol}"
                row[6].text = f"{perf_percent:.2f}%"
                
                for cell in row:
                    set_cell_text_style(cell)
                    
def get_quarter_fields(quarter):
    """Return the appropriate fields for the given quarter"""
    quarter_fields = {
        'first': [('pl1', 'pr1', 'pl1_unit_id', 'pr1_unit_id', 'initial', 'initial_unit_id', 'measure')],
        'second': [('pl2', 'pr2', 'pl2_unit_id', 'pr2_unit_id', 'initial', 'initial_unit_id', 'measure')],
        'third': [('pl3', 'pr3', 'pl3_unit_id', 'pr3_unit_id', 'initial', 'initial_unit_id', 'measure')],
        'fourth': [('pl4', 'pr4', 'pl4_unit_id', 'pr4_unit_id', 'initial', 'initial_unit_id', 'measure')],
        'six': [
            ('pl1', 'pr1', 'pl1_unit_id', 'pr1_unit_id', 'initial', 'initial_unit_id', 'measure'),
            ('pl2', 'pr2', 'pl2_unit_id', 'pr2_unit_id', 'initial', 'initial_unit_id', 'measure')
        ],
        'nine': [
            ('pl1', 'pr1', 'pl1_unit_id', 'pr1_unit_id', 'initial', 'initial_unit_id', 'measure'),
            ('pl2', 'pr2', 'pl2_unit_id', 'pr2_unit_id', 'initial', 'initial_unit_id', 'measure'),
            ('pl3', 'pr3', 'pl3_unit_id', 'pr3_unit_id', 'initial', 'initial_unit_id', 'measure')
        ],
        'year': [
            ('pl1', 'pr1', 'pl1_unit_id', 'pr1_unit_id', 'initial', 'initial_unit_id', 'measure'),
            ('pl2', 'pr2', 'pl2_unit_id', 'pr2_unit_id', 'initial', 'initial_unit_id', 'measure'),
            ('pl3', 'pr3', 'pl3_unit_id', 'pr3_unit_id', 'initial', 'initial_unit_id', 'measure'),
            ('pl4', 'pr4', 'pl4_unit_id', 'pr4_unit_id', 'initial', 'initial_unit_id', 'measure')
        ]
    }
    return quarter_fields.get(quarter or 'year', quarter_fields['year'])

def calculate_kpi_totals(kp, quarter_fields):
    """Calculate total plan and performance for a KPI based on quarter"""
    unit_lookup = {}
    total_plan = 0
    total_perf = 0
    symbol = ""
    init_base_unit = 0
    
    for pl, pr, pl_unit, pr_unit, initial, initial_unit, measure in quarter_fields:
        plan = getattr(kp, pl, 0) or 0
        performance = getattr(kp, pr, 0) or 0
        plan_unit_obj = getattr(kp, pl_unit, None)
        performance_unit_obj = getattr(kp, pr_unit, None)
        initial_value = getattr(kp, initial, 0) or 0
        initial_unit_obj = getattr(kp, initial_unit, None)
        measure_obj = getattr(kp, measure, None)
        
        if measure_obj and measure_obj not in unit_lookup:
            base_unit = Unit.objects.filter(measure_id=measure_obj, isBaseUnit=True).first()
            unit_lookup[measure_obj] = base_unit
        
        base_unit = unit_lookup.get(measure_obj)
        symbol = base_unit.symbol if base_unit and base_unit.symbol else ""
        
        plan_base_unit = plan * plan_unit_obj.conversionFactor if plan_unit_obj else 0
        perf_base_unit = performance * performance_unit_obj.conversionFactor if performance_unit_obj else 0
        init_base_unit = initial_value * initial_unit_obj.conversionFactor if initial_unit_obj else 0
        
        total_plan += plan_base_unit
        total_perf += perf_base_unit
    
    if kp.operation == "average":
        num_quarters = len(quarter_fields)
        total_plan = total_plan / num_quarters if num_quarters > 0 else 0
        total_perf = total_perf / num_quarters if num_quarters > 0 else 0
    
    return total_plan, total_perf, symbol, init_base_unit


def set_cell_bg_color(cell, color):
    """Set background color for a cell"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading_elm.append(shading)
    
def set_cell_text_color(cell, color):
    """Set text color for a cell"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Ebrima'
            run.font.color.rgb = color
            
def set_cell_text_style(cell):
    """Set standard text style for a cell"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Ebrima'
            
@api_view(['GET'])
def filter_and_generate_kpi_report(request):
    """Generate standalone KPI report table"""
    try:
        doc = Document()
        
        filters, error = build_filters(request)
        if error:
            return JsonResponse({"error": error}, status=400)
        
        generate_kpi_performance_table(doc, request, filters)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            temp_docx_path = tmp_file.name
            doc.save(temp_docx_path)
        

        media_dir = os.path.join(settings.MEDIA_ROOT, "documents")
        os.makedirs(media_dir, exist_ok=True)
        
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        user_id = request.user.id
        base_filename = f"kpi_report_{timestamp}_{user_id}"
        docx_filename_final = f"{base_filename}.docx"
        pdf_filename_final = f"{base_filename}.pdf"
        
        final_docx_path = os.path.join(media_dir, docx_filename_final)
        
        with open(temp_docx_path, "rb") as f_read:
            default_storage.save(os.path.join("documents", docx_filename_final), f_read)
        
        try:
            subprocess.run([
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", media_dir,
                final_docx_path
            ], check=True, timeout=60)
        except subprocess.CalledProcessError as e:
            return JsonResponse({"error": f"LibreOffice conversion failed: {e}"}, status=500)
        except subprocess.TimeoutExpired:
            return JsonResponse({"error": "PDF conversion timed out"}, status=500)
        
        os.remove(temp_docx_path)
        
        doc_url_base = request.build_absolute_uri('/media/documents/')
        docx_url = doc_url_base + docx_filename_final
        pdf_url = doc_url_base + pdf_filename_final
        
        return JsonResponse({
            "message": "Document generated successfully",
            "docx_file_path": docx_url,
            "pdf_file_path": pdf_url
        }, status=200)
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({"error": f"An error occurred: {e}"}, status=500)